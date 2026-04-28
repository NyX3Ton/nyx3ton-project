# app.py of nyx3ton-project\CV Validator
from __future__ import annotations

# -----------------------------------------------------------------------------
# 0. IMPORTS
# -----------------------------------------------------------------------------
import gc, json, os, re, tempfile, traceback, torch, requests, faiss
from dataclasses import dataclass
from importlib import import_module
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, cast
import numpy as np
import gradio as gr
from bs4 import BeautifulSoup
from sentence_transformers import SentenceTransformer
from transformers import AutoModelForCausalLM, AutoTokenizer, BitsAndBytesConfig
from dotenv import load_dotenv
# -----------------------------------------------------------------------------
# 1. ENV + GLOBAL SETTINGS
# -----------------------------------------------------------------------------
DEFAULT_LLM_MODEL_ID = os.getenv("LLM_MODEL_ID", "Qwen/Qwen2.5-7B-Instruct")
DEFAULT_FALLBACK_LLM_MODEL_ID = os.getenv("FALLBACK_LLM_MODEL_ID", "Qwen/Qwen2.5-3B-Instruct")
DEFAULT_EMBED_MODEL_ID = os.getenv("EMBED_MODEL_ID","sentence-transformers/paraphrase-multilingual-mpnet-base-v2",)

LLM_LOAD_MODE = os.getenv("LLM_LOAD_MODE", "auto")  # auto | bnb_4bit | fp16_gpu | cpu
MAX_GPU_MEMORY = os.getenv("MAX_GPU_MEMORY", "10.5GiB")
MAX_INPUT_TOKENS = int(os.getenv("MAX_INPUT_TOKENS", "8192"))
MAX_NEW_TOKENS = int(os.getenv("MAX_NEW_TOKENS", "900"))
CHUNK_WORDS = int(os.getenv("CHUNK_WORDS", "220"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "55"))
DEFAULT_TOP_K = int(os.getenv("TOP_K", "20"))
DEFAULT_MAX_REQUIREMENTS = int(os.getenv("MAX_REQUIREMENTS", "12"))
MIN_RAG_SIMILARITY = float(os.getenv("MIN_RAG_SIMILARITY", "0.20"))
TEMPERATURE_SETTING = float(os.getenv("DEF_TEMPERATURE_SETTING", "0.20"))
SAMPLE_SETTING = os.getenv("DEF_SAMPLE_SETTING", "true")
P_SETTING = float(os.getenv("DEF_P_SETTING", "0.20"))

HF_HOME_LOCAL = os.getenv("HF_HOME_LOCAL", "").strip()
if HF_HOME_LOCAL:
    os.environ["HF_HOME"] = str(Path(HF_HOME_LOCAL).expanduser().resolve())

# -----------------------------------------------------------------------------
# 2. GLOBAL MODEL CACHE
# -----------------------------------------------------------------------------

_TOKENIZER = None
_MODEL = None
_MODEL_INFO = "Model este nie je nacitany."
_EMBEDDER = None
_EMBEDDER_ID = None

@dataclass
class Requirement:
    id: str
    text: str
    category: str = "other"
    priority: str = "unknown"
    weight: float = 1.0


# -----------------------------------------------------------------------------
# 3. UTILS
# -----------------------------------------------------------------------------

def require_runtime() -> None:
    if MISSING_IMPORTS:
        unique = sorted(set(MISSING_IMPORTS))
        raise RuntimeError(
            "Chybaju Python kniznice: " + ", ".join(unique) +
            "\n\nNainstaluj ich prikazom:\n    pip install -r requirements.txt"
        )


def safe_json_loads(text: str, fallback: Any) -> Any:
    if not text:
        return fallback

    cleaned = text.strip()
    cleaned = re.sub(r"^```(?:json)?", "", cleaned, flags=re.IGNORECASE).strip()
    cleaned = re.sub(r"```$", "", cleaned).strip()

    # Direct parse
    try:
        return json.loads(cleaned)
    except Exception:
        pass

    # Try object
    obj_start = cleaned.find("{")
    obj_end = cleaned.rfind("}")
    if obj_start != -1 and obj_end != -1 and obj_end > obj_start:
        try:
            return json.loads(cleaned[obj_start:obj_end + 1])
        except Exception:
            pass

    # Try array
    arr_start = cleaned.find("[")
    arr_end = cleaned.rfind("]")
    if arr_start != -1 and arr_end != -1 and arr_end > arr_start:
        try:
            return json.loads(cleaned[arr_start:arr_end + 1])
        except Exception:
            pass

    return fallback

def normalize_space(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()

def trim_text(text: str, max_chars: int = 24000) -> str:
    text = text.strip()
    if len(text) <= max_chars:
        return text
    # Keep beginning and end; CV often has important skills at both places.
    half = max_chars // 2
    return text[:half] + "\n\n[...SKRATENE...]\n\n" + text[-half:]

def file_ext(path: str) -> str:
    return Path(path).suffix.lower().replace(".", "")

def weighted_average(rows: List[Dict[str, Any]]) -> float:
    total_w = 0.0
    total = 0.0
    for r in rows:
        try:
            w = float(r.get("weight", 1.0))
            s = float(r.get("score", 0.0))
        except Exception:
            continue
        total_w += max(w, 0.01)
        total += max(0.0, min(100.0, s)) * max(w, 0.01)
    return round(total / total_w, 2) if total_w else 0.0


# -----------------------------------------------------------------------------
# 4. DOCUMENT LOADING: PDF/DOCX/RTF/TXT/DOC
# -----------------------------------------------------------------------------

def load_pdf(path: str) -> str:
    fitz = import_module("fitz")  # PyMuPDF

    parts = []
    with fitz.open(path) as doc:
        page_count = int(getattr(doc, "page_count", 0))

        for page_index in range(page_count):
            page = doc.load_page(page_index)
            text = page.get_text("text") or ""

            if text.strip():
                parts.append(f"\n--- PAGE {page_index + 1} ---\n{text}")

    return normalize_space("\n".join(parts))

def load_docx(path: str) -> str:
    import docx

    d = docx.Document(path)
    paragraphs = [p.text for p in d.paragraphs if p.text and p.text.strip()]

    # Include tables as text too.
    for table in d.tables:
        for row in table.rows:
            cells = [normalize_space(c.text) for c in row.cells if c.text]
            if cells:
                paragraphs.append(" | ".join(cells))

    return normalize_space("\n".join(paragraphs))

def load_rtf(path: str) -> str:
    from striprtf.striprtf import rtf_to_text

    raw = Path(path).read_text(errors="ignore")
    return normalize_space(rtf_to_text(raw))


def load_txt(path: str) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp1250", "latin-1"):
        try:
            return normalize_space(Path(path).read_text(encoding=enc, errors="ignore"))
        except Exception:
            continue
    return normalize_space(Path(path).read_text(errors="ignore"))

def load_doc_legacy_windows(path: str) -> str:
    try:
        import win32com.client  # type: ignore
    except Exception as exc:
        raise RuntimeError(
            "Subor .doc je legacy format. Pre native Windows fallback treba mat nainstalovany "
            "Microsoft Word + pywin32. Alternativa: uloz CV ako .docx alebo .pdf."
        ) from exc

    tmp_dir = Path(tempfile.mkdtemp(prefix="cvdoc_"))
    tmp_txt = tmp_dir / "converted.txt"
    word = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(Path(path).resolve()))
        # 7 = wdFormatUnicodeText
        doc.SaveAs(str(tmp_txt), FileFormat=7)
        doc.Close(False)
        return load_txt(str(tmp_txt))
    finally:
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass

def load_document(path: str) -> str:
    ext = file_ext(path)
    if ext == "pdf":
        return load_pdf(path)
    if ext == "docx":
        return load_docx(path)
    if ext == "rtf":
        return load_rtf(path)
    if ext in {"txt", "md"}:
        return load_txt(path)
    if ext == "doc":
        return load_doc_legacy_windows(path)
    raise ValueError(f"Nepodporovany format suboru: .{ext}. Pouzi PDF, DOCX, RTF, TXT alebo DOC.")


# -----------------------------------------------------------------------------
# 5. JOB AD SCRAPING
# -----------------------------------------------------------------------------

def scrape_url(url: str) -> str:
    if not url or not url.strip():
        return ""
    if requests is None or BeautifulSoup is None:
        raise RuntimeError("Chyba requests/beautifulsoup4. Spusti: pip install requests beautifulsoup4")

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122 Safari/537.36"
        )
    }
    resp = requests.get(url.strip(), headers=headers, timeout=25)
    resp.raise_for_status()

    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup(["script", "style", "noscript", "svg", "nav", "footer", "header"]):
        tag.decompose()

    # Prefer main/article if available.
    main = soup.find("main") or soup.find("article") or soup.body or soup
    text = main.get_text("\n")
    lines = [normalize_space(x) for x in text.splitlines()]
    lines = [x for x in lines if len(x) > 2]

    # Remove excessive duplicates while preserving order.
    seen = set()
    unique = []
    for line in lines:
        key = line.lower()
        if key in seen:
            continue
        seen.add(key)
        unique.append(line)

    return normalize_space("\n".join(unique))


# -----------------------------------------------------------------------------
# 6. CHUNKING + RAG
# -----------------------------------------------------------------------------

def chunk_text(text: str, words_per_chunk: int = CHUNK_WORDS, overlap: int = CHUNK_OVERLAP) -> List[str]:
    words = text.split()
    if not words:
        return []
    chunks = []
    step = max(1, words_per_chunk - overlap)
    for start in range(0, len(words), step):
        chunk_words = words[start:start + words_per_chunk]
        if len(chunk_words) < 20 and chunks:
            break
        chunks.append(" ".join(chunk_words))
    return chunks


def get_embedder(model_id: str = DEFAULT_EMBED_MODEL_ID):
    global _EMBEDDER, _EMBEDDER_ID
    require_runtime()
    if _EMBEDDER is not None and _EMBEDDER_ID == model_id:
        return _EMBEDDER
    device = "cpu"  # Keep VRAM free for LLM. Change to cuda if you really want.
    _EMBEDDER = SentenceTransformer(model_id, device=device)
    _EMBEDDER_ID = model_id
    return _EMBEDDER


def build_faiss_index(chunks: List[str], embed_model_id: str) -> Tuple[Any, Any]:
    if not chunks:
        raise ValueError("CV neobsahuje ziadny pouzitelny text/chunk.")

    embedder = get_embedder(embed_model_id)

    vectors = embedder.encode(
        chunks,
        convert_to_numpy=True,
        normalize_embeddings=True,
        show_progress_bar=False,
    )

    np_mod = cast(Any, np)
    vectors = np_mod.asarray(vectors, dtype="float32")

    if vectors.ndim != 2:
        raise ValueError(f"Embedding model vratil necakany tvar vektorov: {vectors.shape}")

    faiss_mod = cast(Any, faiss)
    index = cast(Any, faiss_mod.IndexFlatIP(int(vectors.shape[1])))

    add_fn = cast(Any, index.add)
    add_fn(vectors)

    return index, vectors


def rag_search(query: str, chunks: List[str], index: Any, embed_model_id: str, top_k: int) -> List[str]:
    embedder = get_embedder(embed_model_id)
    q = embedder.encode([query], convert_to_numpy=True, normalize_embeddings=True, show_progress_bar=False).astype("float32")
    scores, ids = index.search(q, min(top_k, len(chunks)))
    results = []
    for score, idx in zip(scores[0], ids[0]):
        if idx < 0:
            continue
        if float(score) < MIN_RAG_SIMILARITY:
            continue
        results.append(f"[similarity={float(score):.3f}] {chunks[int(idx)]}")
    return results


# -----------------------------------------------------------------------------
# 7. LOCAL HUGGING FACE LLM
# -----------------------------------------------------------------------------

def cuda_summary() -> str:
    if torch is None:
        return "PyTorch nie je dostupny."
    if not torch.cuda.is_available():
        return "CUDA nie je dostupna, pouzije sa CPU alebo fallback."
    name = torch.cuda.get_device_name(0)
    total_gb = torch.cuda.get_device_properties(0).total_memory / 1024 ** 3
    allocated_gb = torch.cuda.memory_allocated(0) / 1024 ** 3
    reserved_gb = torch.cuda.memory_reserved(0) / 1024 ** 3
    return f"CUDA OK: {name}, VRAM total={total_gb:.1f} GB, allocated={allocated_gb:.2f} GB, reserved={reserved_gb:.2f} GB"


def unload_llm() -> str:
    global _TOKENIZER, _MODEL, _MODEL_INFO
    _TOKENIZER = None
    _MODEL = None
    _MODEL_INFO = "Model bol uvolneny."
    gc.collect()
    if torch is not None and torch.cuda.is_available():
        torch.cuda.empty_cache()
    return _MODEL_INFO + "\n" + cuda_summary()


def load_llm(
    model_id: str = DEFAULT_LLM_MODEL_ID,
    load_mode: str = LLM_LOAD_MODE,
    fallback_model_id: str = DEFAULT_FALLBACK_LLM_MODEL_ID,
):
    global _TOKENIZER, _MODEL, _MODEL_INFO
    require_runtime()

    desired_signature = f"{model_id}|{load_mode}|fallback={fallback_model_id}"
    if _MODEL is not None and _TOKENIZER is not None and desired_signature in _MODEL_INFO:
        return _TOKENIZER, _MODEL, _MODEL_INFO

    unload_llm()

    errors = []
    has_cuda = torch.cuda.is_available()

    def _load_tokenizer(mid: str):
        return AutoTokenizer.from_pretrained(mid, trust_remote_code=True)

    def _try_4bit(mid: str):
        if not has_cuda:
            raise RuntimeError("CUDA nie je dostupna pre 4-bit GPU load.")
        bnb_config = BitsAndBytesConfig(
                                        load_in_4bit=True,
            bnb_4bit_quant_type="nf4",
            bnb_4bit_compute_dtype=torch.float16,
            bnb_4bit_use_double_quant=True,
        )
        tok = _load_tokenizer(mid)
        mdl = AutoModelForCausalLM.from_pretrained(
            mid,
            quantization_config=bnb_config,
            device_map="auto",
            max_memory={0: MAX_GPU_MEMORY, "cpu": "32GiB"},
            trust_remote_code=True,
        )
        return tok, mdl, f"Nacitany model: {mid} | mode=bnb_4bit | {desired_signature}"

    def _try_fp16_gpu(mid: str):
        if not has_cuda:
            raise RuntimeError("CUDA nie je dostupna pre fp16_gpu load.")
        tok = _load_tokenizer(mid)
        mdl = AutoModelForCausalLM.from_pretrained(
            mid,
            torch_dtype=torch.float16,
            device_map="auto",
            max_memory={0: MAX_GPU_MEMORY, "cpu": "32GiB"},
            trust_remote_code=True,
        )
        return tok, mdl, f"Nacitany model: {mid} | mode=fp16_gpu | {desired_signature}"

    def _try_cpu(mid: str):
        tok = _load_tokenizer(mid)
        mdl = AutoModelForCausalLM.from_pretrained(
            mid,
            torch_dtype=torch.float32,
            device_map={"": "cpu"},
            trust_remote_code=True,
        )
        return tok, mdl, f"Nacitany model: {mid} | mode=cpu | {desired_signature}"

    attempts = []
    mode = (load_mode or "auto").lower().strip()

    if mode == "bnb_4bit":
        attempts = [lambda: _try_4bit(model_id)]
    elif mode == "fp16_gpu":
        attempts = [lambda: _try_fp16_gpu(model_id)]
    elif mode == "cpu":
        attempts = [lambda: _try_cpu(model_id)]
    else:
        attempts = [
            lambda: _try_4bit(model_id),
            lambda: _try_fp16_gpu(fallback_model_id),
            lambda: _try_cpu(fallback_model_id),
        ]

    for attempt in attempts:
        try:
            _TOKENIZER, _MODEL, _MODEL_INFO = attempt()
            _MODEL.eval()
            return _TOKENIZER, _MODEL, _MODEL_INFO + "\n" + cuda_summary()
        except Exception as exc:
            errors.append(f"{type(exc).__name__}: {exc}")
            gc.collect()
            if torch.cuda.is_available():
                torch.cuda.empty_cache()

    raise RuntimeError(
        "Nepodarilo sa nacitat lokalny LLM.\n\nPokusy:\n- " + "\n- ".join(errors)
    )

def model_device(model: Any):
    try:
        return next(model.parameters()).device
    except Exception:
        return "cuda" if torch.cuda.is_available() else "cpu"

def chat_generate(
    system: str,
    user: str,
    model_id: str,
    load_mode: str,
    fallback_model_id: str,
    max_new_tokens: int = MAX_NEW_TOKENS,
) -> str:
    tok, mdl, _ = load_llm(model_id, load_mode, fallback_model_id)

    messages = [{"role": "system", "content": system.strip()},{"role": "user", "content": user.strip()}]

    if getattr(tok, "chat_template", None):
        prompt = tok.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)
    else:
        prompt = f"System:\n{system.strip()}\n\nUser:\n{user.strip()}\n\nAssistant:\n"

    inputs = tok(prompt,return_tensors="pt",truncation=True,max_length=MAX_INPUT_TOKENS)
    dev = model_device(mdl)
    inputs = {k: v.to(dev) for k, v in inputs.items()}

    with torch.no_grad():
        out = mdl.generate(
            **inputs,
            max_new_tokens=max_new_tokens,
            temperature=TEMPERATURE_SETTING,
            do_sample=SAMPLE_SETTING,
            top_p=P_SETTING,
            repetition_penalty=1.05,
            pad_token_id=tok.eos_token_id,
        )

    # Decode only newly generated part where possible.
    generated = out[0][inputs["input_ids"].shape[-1]:]
    text = tok.decode(generated, skip_special_tokens=True)
    return text.strip()

SYSTEM_JSON = """
                Si lokalny AI asistent pre validaciu zivotopisov voci pracovnemu inzeratu.

                Tvoje pravidla:
                - odpovedaj iba validnym JSON objektom alebo validnym JSON polom
                - nepouzivaj markdown
                - nepouzivaj text mimo JSON
                - nehadaj informacie, ktore nie su v dodanom texte
                - pri hodnoteni kandidata pouzivaj iba dodane dokazy z CV
                - nevykonavaj finalne rozhodnutie o prijati kandidata
                - vystup sluzi iba ako odporucanie pre cloveka
                - nehodnot citlive atributy ako vek, pohlavie, narodnost, zdravotny stav, rodinny stav, fotografia alebo adresa
                - pis slovensky bez diakritiky
                """.strip()
                
# -----------------------------------------------------------------------------
# 8. LLM PROMPTS / TASKS
# -----------------------------------------------------------------------------
def extract_job_requirements(
    job_text: str,
    model_id: str,
    load_mode: str,
    fallback_model_id: str,
    max_requirements: int,
) -> Dict[str, Any]:
    schema = """
{
    "job_title": "",
    "seniority": "junior|medior|senior|unknown",
    "requirements": [
        {
        "id": "R1",
        "text": "konkretna poziadavka",
        "category": "hard_skill|experience|education|language|soft_skill|location|other",
        "priority": "must|nice|unknown",
        "weight": 1.0
        }]
}
""".strip()
    user = (
            "Z pracovnej ponuky extrahuj atomicke poziadavky na kandidata.\n\n"
            "Vrat presne tento JSON tvar:\n" + schema + "\n\n"
            "Pravidla:\n"
            f"- maximalne {max_requirements} najdolezitejsich poziadaviek\n"
            "- must-have poziadavky daj weight 3 az 5\n"
            "- nice-to-have poziadavky daj weight 1 az 2\n"
            "- nerozbijaj jednu poziadavku na duplicitne varianty\n"
            "- ignoruj benefity firmy, marketing a pravne formulacie\n\n"
            "PRACOVNA PONUKA:\n" + trim_text(job_text, 22000)
            )
    raw = chat_generate(SYSTEM_JSON, user, model_id, load_mode, fallback_model_id, max_new_tokens=1200)
    data = safe_json_loads(raw, fallback={"job_title": "unknown", "seniority": "unknown", "requirements": []})

    # Basic normalization in case model returns a list directly.
    if isinstance(data, list):
        data = {"job_title": "unknown", "seniority": "unknown", "requirements": data}
    if not isinstance(data, dict):
        data = {"job_title": "unknown", "seniority": "unknown", "requirements": []}

    reqs = data.get("requirements") or []
    clean_reqs = []
    for i, r in enumerate(reqs[:max_requirements], start=1):
        if not isinstance(r, dict):
            continue
        text = normalize_space(str(r.get("text", "")))
        if not text:
            continue
        try:
            weight = float(r.get("weight", 1.0))
        except Exception:
            weight = 1.0
        clean_reqs.append({
            "id": str(r.get("id") or f"R{i}"),
            "text": text,
            "category": str(r.get("category") or "other"),
            "priority": str(r.get("priority") or "unknown"),
            "weight": max(0.5, min(5.0, weight)),
        })

    data["requirements"] = clean_reqs
    return data

def extract_candidate_summary(
    cv_text: str,
    model_id: str,
    load_mode: str,
    fallback_model_id: str,
) -> Dict[str, Any]:
    schema = """
    {
        "summary": "kratke zhrnutie kandidata bez mena a citlivych atributov",
        "skills": ["..."],
        "experience": ["..."],
        "education": ["..."],
        "languages": ["..."],
        "certifications": ["..."],
        "risks_or_missing_info": ["..."]
    }
""".strip()
    user = (
                "Zo zivotopisu extrahuj anonymizovany profil kandidata.\n\n"
                "Vrat presne tento JSON tvar:\n" + schema + "\n\n"
                "Pravidla:\n"
                "- nepouzivaj meno, adresu, vek, pohlavie, rodinny stav ani fotku\n"
                "- uvadzaj len veci, ktore su v CV\n\n"
                "CV TEXT:\n" + trim_text(cv_text, 22000)
            )
    raw = chat_generate(SYSTEM_JSON, user, model_id, load_mode, fallback_model_id, max_new_tokens=1100)
    data = safe_json_loads(raw, fallback={"summary": "Nepodarilo sa spolahlivo extrahovat profil.", "skills": []})
    if not isinstance(data, dict):
        data = {"summary": "Nepodarilo sa spolahlivo extrahovat profil.", "skills": []}
    return data

def evaluate_one_requirement(
    requirement: Dict[str, Any],
    evidence: List[str],
    model_id: str,
    load_mode: str,
    fallback_model_id: str,
) -> Dict[str, Any]:
    schema = {
                "requirement_id": str(requirement.get("id", "")),
                "requirement": str(requirement.get("text", "")),
                "status": "splnene|ciastocne_splnene|nesplnene|nejasne",
                "score": 0,
                "confidence": 0.0,
                "evidence_used": ["kratke citacie/parafrazy dokazov"],
                "explanation": "kratke vysvetlenie v slovencine bez diakritiky",
                "risk_note": "co chyba alebo preco je hodnotenie neiste",
                }
    user = (
            "Vyhodnot, ci kandidat splna jednu poziadavku z pracovneho inzeratu.\n"
            "Pouzi iba dodane dokazy z CV. Ak dokaz chyba, nehadaj.\n\n"
            "Vrat presne tento JSON tvar:\n"
            + json.dumps(schema, ensure_ascii=False, indent=2)
            + "\n\nSkorovanie:\n"
            "- 90-100 = jasne splnene\n"
            "- 60-89 = skor splnene alebo ciastocne\n"
            "- 30-59 = slabe/nepriame dokazy\n"
            "- 0-29 = nesplnene alebo chyba dokaz\n\n"
            "POZIADAVKA:\n" + json.dumps(requirement, ensure_ascii=False) + "\n\n"
            "DOKAZY Z CV:\n" + json.dumps(evidence, ensure_ascii=False, indent=2)
            )
    raw = chat_generate(SYSTEM_JSON, user, model_id, load_mode, fallback_model_id, max_new_tokens=850)
    data = safe_json_loads(raw, fallback={})
    if not isinstance(data, dict):
        data = {}

    # Normalize robustly.
    status = str(data.get("status") or "nejasne")
    if status not in {"splnene", "ciastocne_splnene", "nesplnene", "nejasne"}:
        status = "nejasne"

    try:
        score = float(data.get("score", 0))
    except Exception:
        score = 0.0
    try:
        confidence = float(data.get("confidence", 0))
    except Exception:
        confidence = 0.0

    return {
            "requirement_id": str(data.get("requirement_id") or requirement.get("id", "")),
            "requirement": str(data.get("requirement") or requirement.get("text", "")),
            "category": requirement.get("category", "other"),
            "priority": requirement.get("priority", "unknown"),
            "weight": requirement.get("weight", 1.0),
            "status": status,
            "score": max(0.0, min(100.0, score)),
            "confidence": max(0.0, min(1.0, confidence)),
            "evidence_used": data.get("evidence_used") if isinstance(data.get("evidence_used"), list) else evidence[:2],
            "explanation": str(data.get("explanation") or "Bez vysvetlenia."),
            "risk_note": str(data.get("risk_note") or ""),
            }

# -----------------------------------------------------------------------------
# 9. REPORTING
# -----------------------------------------------------------------------------

def status_icon(status: str) -> str:
    return {
            "splnene": "✅",
            "ciastocne_splnene": "🟡",
            "nesplnene": "❌",
            "nejasne": "⚪",
            }.get(status, "⚪")


def verdict(score: float) -> str:
    if score >= 80:
        return "Vhodny kandidat"
    if score >= 60:
        return "Skor vhodny kandidat"
    if score >= 40:
        return "Ciastocne vhodny / vyzaduje manualne posudenie"
    return "Slaba zhoda s poziciou"


def render_markdown_report(job_data: Dict[str, Any], candidate: Dict[str, Any], evals: List[Dict[str, Any]]) -> str:
    score = weighted_average(evals)
    lines = []
    lines.append(f"# Lokalny AI CV validator")
    lines.append("")
    lines.append(f"**Pozicia:** {job_data.get('job_title', 'unknown')}")
    lines.append(f"**Seniorita:** {job_data.get('seniority', 'unknown')}")
    lines.append(f"**Celkove skore:** {score:.2f} / 100")
    lines.append(f"**Odporucanie:** {verdict(score)}")
    lines.append("")
    lines.append("> Vystup je odporucanie pre cloveka, nie automaticke rozhodnutie o kandidatovi.")
    lines.append("")

    if candidate:
        lines.append("## Anonymizovany profil kandidata")
        lines.append(str(candidate.get("summary", "")))
        for key, title in [
            ("skills", "Zrucnosti"),
            ("experience", "Skusenosti"),
            ("education", "Vzdelanie"),
            ("languages", "Jazyky"),
            ("certifications", "Certifikaty"),
            ("risks_or_missing_info", "Rizika / chybajuce info"),
        ]:
            vals = candidate.get(key)
            if isinstance(vals, list) and vals:
                lines.append(f"\n**{title}:**")
                for v in vals[:12]:
                    lines.append(f"- {v}")
        lines.append("")

    lines.append("## Vyhodnotenie poziadaviek")
    lines.append("")
    lines.append("| Stav | Poziadavka | Skore | Priorita | Vysvetlenie |")
    lines.append("|---|---|---:|---|---|")
    for r in evals:
        exp = str(r.get("explanation", "")).replace("|", "/")
        req = str(r.get("requirement", "")).replace("|", "/")
        lines.append(
            f"| {status_icon(r.get('status', ''))} {r.get('status', '')} "
            f"| {req} | {float(r.get('score', 0)):.0f} "
            f"| {r.get('priority', '')} / w={r.get('weight', 1)} | {exp} |"
        )

    lines.append("")
    lines.append("## Dokazy a poznamky")
    for r in evals:
        lines.append(f"\n### {status_icon(r.get('status', ''))} {r.get('requirement_id', '')}: {r.get('requirement', '')}")
        if r.get("risk_note"):
            lines.append(f"**Riziko/neistota:** {r.get('risk_note')}")
        ev = r.get("evidence_used") or []
        if isinstance(ev, list) and ev:
            lines.append("**Pouzite dokazy:**")
            for e in ev[:3]:
                short = normalize_space(str(e))[:700]
                lines.append(f"- {short}")

    return "\n".join(lines)


# -----------------------------------------------------------------------------
# 10. MAIN PIPELINE
# -----------------------------------------------------------------------------

def run_validation(
    cv_file: str,
    job_url: str,
    job_text_manual: str,
    model_id: str,
    fallback_model_id: str,
    load_mode: str,
    embed_model_id: str,
    top_k: int,
    max_requirements: int,
    include_candidate_summary: bool,
) -> Tuple[str, str, str]:
    """Returns: markdown_report, json_report, runtime_info"""
    require_runtime()

    if not cv_file:
        raise gr.Error("Nahraj CV subor.")

    runtime = []
    runtime.append(cuda_summary())
    runtime.append(f"LLM model: {model_id}")
    runtime.append(f"Fallback model: {fallback_model_id}")
    runtime.append(f"Load mode: {load_mode}")
    runtime.append(f"Embedding model: {embed_model_id}")

    cv_text = load_document(cv_file)
    if len(cv_text) < 100:
        raise gr.Error("Z CV sa podarilo vytiahnut velmi malo textu. Skus iny format, idealne PDF/DOCX.")
    runtime.append(f"CV text: {len(cv_text):,} znakov")

    job_text = ""
    if job_text_manual and job_text_manual.strip():
        job_text = normalize_space(job_text_manual)
        runtime.append("Inzerat: pouzity manualne vlozeny text")
    elif job_url and job_url.strip():
        job_text = scrape_url(job_url)
        runtime.append(f"Inzerat: nacitany z URL, {len(job_text):,} znakov")
    else:
        raise gr.Error("Zadaj URL inzeratu alebo vloz text inzeratu manualne.")

    if len(job_text) < 100:
        raise gr.Error("Z inzeratu sa podarilo ziskat velmi malo textu. Vloz text inzeratu manualne.")

    # Load LLM once before subtasks.
    _, _, model_info = load_llm(model_id, load_mode, fallback_model_id)
    runtime.append(model_info)

    job_data = extract_job_requirements(job_text, model_id, load_mode, fallback_model_id, max_requirements)
    requirements = job_data.get("requirements", [])
    if not requirements:
        raise gr.Error("LLM neextrahoval ziadne poziadavky z inzeratu. Skus vlozit cistejsi text inzeratu.")
    runtime.append(f"Extrahovane poziadavky: {len(requirements)}")

    candidate = {}
    if include_candidate_summary:
        candidate = extract_candidate_summary(cv_text, model_id, load_mode, fallback_model_id)

    chunks = chunk_text(cv_text, CHUNK_WORDS, CHUNK_OVERLAP)
    index, _ = build_faiss_index(chunks, embed_model_id)
    runtime.append(f"CV chunks: {len(chunks)}")

    evals = []
    for req in requirements:
        evidence = rag_search(req.get("text", ""), chunks, index, embed_model_id, int(top_k))
        ev = evaluate_one_requirement(req, evidence, model_id, load_mode, fallback_model_id)
        evals.append(ev)

    final = {
        "job": job_data,
        "candidate_profile": candidate,
        "overall_score": weighted_average(evals),
        "verdict": verdict(weighted_average(evals)),
        "evaluations": evals,
        "runtime": runtime,
    }

    md = render_markdown_report(job_data, candidate, evals)
    js = json.dumps(final, ensure_ascii=False, indent=2)
    return md, js, "\n".join(runtime)


def gradio_run_wrapper(*args):
    try:
        return run_validation(*args)
    except gr.Error:
        raise
    except Exception as exc:
        err = f"Chyba: {type(exc).__name__}: {exc}\n\n{traceback.format_exc()}"
        return "# Chyba pri spracovani\n\n```text\n" + err + "\n```", "{}", err


# -----------------------------------------------------------------------------
# 11. GRADIO UI
# -----------------------------------------------------------------------------

def build_ui():
    require_runtime()

    with gr.Blocks(title="Lokalny AI CV Validator") as demo:
        gr.Markdown(
        )

        with gr.Row():
            with gr.Column(scale=1):
                cv_file = gr.File(
                    label="CV subor",
                    file_types=[".pdf", ".docx", ".doc", ".rtf", ".txt", ".md"],
                    type="filepath",
                )
                job_url = gr.Textbox(label="URL inzeratu", placeholder="https://...")
                job_text_manual = gr.Textbox(
                    label="Alebo vloz text inzeratu manualne",
                    lines=8,
                    placeholder="Text pracovnej ponuky",
                )

                with gr.Accordion("Model nastavenia", open=False):
                    model_id = gr.Textbox(label="LLM model z Hugging Face", value=DEFAULT_LLM_MODEL_ID)
                    fallback_model_id = gr.Textbox(label="Fallback LLM model", value=DEFAULT_FALLBACK_LLM_MODEL_ID)
                    load_mode = gr.Dropdown(
                        label="Load mode",
                        choices=["auto", "bnb_4bit", "fp16_gpu", "cpu"],
                        value=LLM_LOAD_MODE,
                    )
                    embed_model_id = gr.Textbox(label="Embedding model z Hugging Face", value=DEFAULT_EMBED_MODEL_ID)
                    top_k = gr.Slider(label="Top-K dokazov z CV", minimum=2, maximum=10, step=1, value=DEFAULT_TOP_K)
                    max_requirements = gr.Slider(
                        label="Max pocet poziadaviek z inzeratu",
                        minimum=5,
                        maximum=25,
                        step=1,
                        value=DEFAULT_MAX_REQUIREMENTS,
                    )
                    include_candidate_summary = gr.Checkbox(label="Extrahovat anonymizovany profil kandidata", value=True)

                with gr.Row():
                    run_btn = gr.Button("Spustit validaciu", variant="primary")
                    unload_btn = gr.Button("Uvolnit model z VRAM")

            with gr.Column(scale=2):
                report_md = gr.Markdown(label="Report")
                runtime_info = gr.Textbox(label="Runtime info", lines=8)
                json_report = gr.Code(label="JSON report", language="json", lines=20)

        run_btn.click(
            fn=gradio_run_wrapper,
            inputs=[
                    cv_file,
                    job_url,
                    job_text_manual,
                    model_id,
                    fallback_model_id,
                    load_mode,
                    embed_model_id,
                    top_k,
                    max_requirements,
                    include_candidate_summary,
                    ],
            outputs=[report_md, json_report, runtime_info],
        )
        unload_btn.click(fn=unload_llm, inputs=[], outputs=[runtime_info])

    return demo

if __name__ == "__main__":
    if MISSING_IMPORTS:
        print("Chybaju kniznice:", ", ".join(sorted(set(MISSING_IMPORTS))))
        print("Spusti: pip install -r requirements.txt")
    demo = build_ui()
    demo.launch(server_name="127.0.0.1", server_port=7860, inbrowser=True)
