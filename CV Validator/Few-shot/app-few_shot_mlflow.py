# main script: app-few_shot.py of nyx3ton-project\CV Validator

from __future__ import annotations

# -----------------------------------------------------------------------------
# 0. IMPORTS
# 1. GLOBAL MODEL CACHE
# 2. EXTERNAL HELPERS
# 3. DOCUMENT LOADING: PDF/DOCX/RTF/TXT/DOC
# 4. JOB AD SCRAPING
# 5. CHUNKING + RAG
# 6. REPORTING
# 7. MAIN PIPELINE
# 8. GRADIO UI
# -----------------------------------------------------------------------------

# -----------------------------------------------------------------------------
# 0. IMPORTS
# -----------------------------------------------------------------------------
import gc, json, os, re, sys, socket, atexit, subprocess, webbrowser, tempfile, traceback, time, hashlib, torch, requests, faiss
from importlib import import_module
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, cast
import numpy as np
import gradio as gr
from bs4 import BeautifulSoup
from sentence_transformers import SentenceTransformer
from dotenv import load_dotenv
from dictionary_fallback import (fallback_extract_requirements_from_text, build_hybrid_requirement_result)

# Optional experiment tracking.
# The app still works without MLflow installed; logging is enabled only when available.
from typing import Any

try:
    import mlflow as _mlflow
    MLFLOW_AVAILABLE = True
except ImportError:
    _mlflow = None
    MLFLOW_AVAILABLE = False

mlflow: Any = _mlflow


# -----------------------------------------------------------------------------
# 1. ENV + GLOBAL SETTINGS
# -----------------------------------------------------------------------------
APP_DIR = Path(__file__).resolve().parent
load_dotenv(APP_DIR / ".env")

def env_bool(name: str, default: bool = False) -> bool:
    value = os.getenv(name)
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "y", "on"}

DEFAULT_EMBED_MODEL_ID = os.getenv("EMBED_MODEL_ID","sentence-transformers/paraphrase-multilingual-mpnet-base-v2",)
CHUNK_WORDS = int(os.getenv("CHUNK_WORDS", "220"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "55"))
DEFAULT_TOP_K = max(2, min(20, int(os.getenv("TOP_K", "5"))))
DEFAULT_MAX_REQUIREMENTS = max(5, min(25, int(os.getenv("MAX_REQUIREMENTS", "12"))))
MIN_RAG_SIMILARITY = float(os.getenv("MIN_RAG_SIMILARITY", "0.20"))

HF_HOME_LOCAL = os.getenv("HF_HOME_LOCAL", "").strip()
if HF_HOME_LOCAL:
    os.environ["HF_HOME"] = str(Path(HF_HOME_LOCAL).expanduser().resolve())


# -----------------------------------------------------------------------------
# 1A. MLFLOW EXPERIMENT TRACKING SETTINGS
# -----------------------------------------------------------------------------
MLFLOW_ROOT_DIR = Path(os.getenv("MLFLOW_ROOT_DIR", str(APP_DIR / "mlflow_runs"))).expanduser().resolve()

MLFLOW_DB_PATH = MLFLOW_ROOT_DIR / "mlflow.db"
MLFLOW_ARTIFACTS_DIR = MLFLOW_ROOT_DIR / "artifacts"
MLFLOW_EXPERIMENT_NAME = os.getenv("MLFLOW_EXPERIMENT_NAME","few-shot-cv-validator")
MLFLOW_TOKENIZER_LOCAL_ONLY = env_bool("MLFLOW_TOKENIZER_LOCAL_ONLY", True)

def setup_mlflow() -> bool:
    global MLFLOW_AVAILABLE

    if not MLFLOW_AVAILABLE or mlflow is None:
        print("MLflow is not installed. Experiment tracking is disabled.")
        return False

    try:
        MLFLOW_ROOT_DIR.mkdir(parents=True, exist_ok=True)
        MLFLOW_ARTIFACTS_DIR.mkdir(parents=True, exist_ok=True)
        tracking_uri = os.getenv("MLFLOW_TRACKING_URI", "").strip()

        if not tracking_uri:
            tracking_uri = f"sqlite:///{MLFLOW_DB_PATH.as_posix()}"

        mlflow.set_tracking_uri(tracking_uri)

        existing_experiment = mlflow.get_experiment_by_name(MLFLOW_EXPERIMENT_NAME)

        if existing_experiment is None:
            mlflow.create_experiment(name=MLFLOW_EXPERIMENT_NAME,artifact_location=MLFLOW_ARTIFACTS_DIR.as_uri())

        mlflow.set_experiment(MLFLOW_EXPERIMENT_NAME)

        print("MLflow tracking initialized.")
        print(f"MLflow experiment: {MLFLOW_EXPERIMENT_NAME}")
        print(f"MLflow tracking URI: {tracking_uri}")
        print(f"MLflow artifacts: {MLFLOW_ARTIFACTS_DIR}")

        return True

    except Exception as exc:
        MLFLOW_AVAILABLE = False

        print("MLflow initialization failed. Experiment tracking is disabled.")
        print(f"Reason: {type(exc).__name__}: {exc}")

        return False

MLFLOW_READY = setup_mlflow()
MLFLOW_SETUP_ERROR = ""
START_MLFLOW_UI = env_bool("START_MLFLOW_UI", True)
OPEN_MLFLOW_UI_BROWSER = env_bool("OPEN_MLFLOW_UI_BROWSER", True)
MLFLOW_UI_HOST = os.getenv("MLFLOW_UI_HOST", "127.0.0.1")
MLFLOW_UI_PORT = int(os.getenv("MLFLOW_UI_PORT", "5001"))
_MLFLOW_UI_PROCESS = None


def _is_port_open(host: str, port: int) -> bool:
    try:
        with socket.create_connection((host, port), timeout=1):
            return True
    except OSError:
        return False


def start_mlflow_ui(open_browser: bool = False) -> Optional[subprocess.Popen]:
    """Start MLflow UI as a background process when requested by env variable."""
    global _MLFLOW_UI_PROCESS

    if not MLFLOW_READY:
        print(f"MLflow UI not started: {MLFLOW_SETUP_ERROR or 'MLflow is not ready.'}")
        return None

    if _is_port_open(MLFLOW_UI_HOST, MLFLOW_UI_PORT):
        print(f"MLflow UI already running at http://{MLFLOW_UI_HOST}:{MLFLOW_UI_PORT}")
        if open_browser:
            webbrowser.open(f"http://{MLFLOW_UI_HOST}:{MLFLOW_UI_PORT}")
        return None

    backend_store_uri = f"sqlite:///{MLFLOW_DB_PATH.as_posix()}"

    cmd = [
        sys.executable,
        "-m",
        "mlflow",
        "ui",
        "--backend-store-uri",
        backend_store_uri,
        "--host",
        MLFLOW_UI_HOST,
        "--port",
        str(MLFLOW_UI_PORT),
    ]

    creationflags = 0
    if os.name == "nt":
        creationflags = subprocess.CREATE_NO_WINDOW

    _MLFLOW_UI_PROCESS = subprocess.Popen(
        cmd,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.STDOUT,
        creationflags=creationflags,
    )

    print(f"MLflow UI started at http://{MLFLOW_UI_HOST}:{MLFLOW_UI_PORT}")

    if open_browser:
        time.sleep(2)
        webbrowser.open(f"http://{MLFLOW_UI_HOST}:{MLFLOW_UI_PORT}")

    return _MLFLOW_UI_PROCESS


def stop_mlflow_ui() -> None:
    global _MLFLOW_UI_PROCESS

    if _MLFLOW_UI_PROCESS is not None and _MLFLOW_UI_PROCESS.poll() is None:
        try:
            _MLFLOW_UI_PROCESS.terminate()
        except Exception:
            pass


atexit.register(stop_mlflow_ui)

if START_MLFLOW_UI:
    start_mlflow_ui(open_browser=OPEN_MLFLOW_UI_BROWSER)


# -----------------------------------------------------------------------------
# 1. GLOBAL MODEL CACHE
# -----------------------------------------------------------------------------
_EMBEDDER = None
_EMBEDDER_ID = None

# -----------------------------------------------------------------------------
# 2. EXTERNAL HELPERS
# -----------------------------------------------------------------------------
from validator_utils import file_ext, normalize_space, weighted_average
from validator_llm import (DEFAULT_FALLBACK_LLM_MODEL_ID,DEFAULT_LLM_MODEL_ID,LLM_LOAD_MODE,cuda_summary,load_llm,unload_llm)
from validator_tasks import (DEFAULT_AUX_LLM_MODEL_ID,DEFAULT_JOB_SCHEMA_XLSX_PATH,extract_candidate_summary,extract_job_requirements,evaluate_one_requirement,load_job_requirement_schema_text,load_manual_job_requirements_from_excel)

# -----------------------------------------------------------------------------
# 3. DOCUMENT LOADING: PDF/DOCX/RTF/TXT/DOC
# -----------------------------------------------------------------------------

def load_pdf(path: str) -> str:
    fitz = import_module("fitz")  # PyMuPDF

    parts = []
    with fitz.open(path) as doc:
        page_count = int(getattr(doc, "page_count", 0))

        for page_index in range(page_count):
            page = doc.load_page(page_index)
            text = page.get_text("text") or ""

            if text.strip():
                parts.append(f"\n--- PAGE {page_index + 1} ---\n{text}")

    return normalize_space("\n".join(parts))

def load_docx(path: str) -> str:
    import docx

    d = docx.Document(path)
    paragraphs = [p.text for p in d.paragraphs if p.text and p.text.strip()]

    for table in d.tables:
        for row in table.rows:
            cells = [normalize_space(c.text) for c in row.cells if c.text]
            if cells:
                paragraphs.append(" | ".join(cells))

    return normalize_space("\n".join(paragraphs))

def load_rtf(path: str) -> str:
    from striprtf.striprtf import rtf_to_text

    raw = Path(path).read_text(errors="ignore")
    return normalize_space(rtf_to_text(raw))

def load_txt(path: str) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp1250", "latin-1"):
        try:
            return normalize_space(Path(path).read_text(encoding=enc, errors="ignore"))
        except Exception:
            continue
    return normalize_space(Path(path).read_text(errors="ignore"))

def load_doc_legacy_windows(path: str) -> str:
    try:
        import win32com.client  # type: ignore
    except Exception as exc:
        raise RuntimeError(
            "Subor .doc je legacy format. Pre native Windows fallback treba mat nainstalovany "
            "Microsoft Word + pywin32. Alternativa: uloz CV ako .docx alebo .pdf."
        ) from exc

    tmp_dir = Path(tempfile.mkdtemp(prefix="cvdoc_"))
    tmp_txt = tmp_dir / "converted.txt"
    word = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(Path(path).resolve()))
        doc.SaveAs(str(tmp_txt), FileFormat=7)  # 7 = wdFormatUnicodeText
        doc.Close(False)
        return load_txt(str(tmp_txt))
    finally:
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass

def load_document(path: str) -> str:
    ext = file_ext(path)
    if ext == "pdf":
        return load_pdf(path)
    if ext == "docx":
        return load_docx(path)
    if ext == "rtf":
        return load_rtf(path)
    if ext in {"txt", "md"}:
        return load_txt(path)
    if ext == "doc":
        return load_doc_legacy_windows(path)
    raise ValueError(f"Nepodporovany format suboru: .{ext}. Pouzi PDF, DOCX, RTF, TXT alebo DOC.")

# -----------------------------------------------------------------------------
# 4. JOB AD SCRAPING
# -----------------------------------------------------------------------------

def scrape_url(url: str) -> str:
    if not url or not url.strip():
        return ""

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122 Safari/537.36"
        )
    }
    resp = requests.get(url.strip(), headers=headers, timeout=25)
    resp.raise_for_status()

    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup(["script", "style", "noscript", "svg", "nav", "footer", "header"]):
        tag.decompose()

    main = soup.find("main") or soup.find("article") or soup.body or soup
    text = main.get_text("\n")
    lines = [normalize_space(x) for x in text.splitlines()]
    lines = [x for x in lines if len(x) > 2]

    seen = set()
    unique = []
    for line in lines:
        key = line.lower()
        if key in seen:
            continue
        seen.add(key)
        unique.append(line)

    return normalize_space("\n".join(unique))

# -----------------------------------------------------------------------------
# 5. CHUNKING + RAG
# -----------------------------------------------------------------------------

def chunk_text(text: str, words_per_chunk: int = CHUNK_WORDS, overlap: int = CHUNK_OVERLAP) -> List[str]:
    words = text.split()
    if not words:
        return []
    chunks = []
    step = max(1, words_per_chunk - overlap)
    for start in range(0, len(words), step):
        chunk_words = words[start:start + words_per_chunk]
        if len(chunk_words) < 20 and chunks:
            break
        chunks.append(" ".join(chunk_words))
    return chunks

def get_embedder(model_id: str = DEFAULT_EMBED_MODEL_ID):
    global _EMBEDDER, _EMBEDDER_ID
    if _EMBEDDER is not None and _EMBEDDER_ID == model_id:
        return _EMBEDDER
    device = "cpu"
    _EMBEDDER = SentenceTransformer(model_id, device=device)
    _EMBEDDER_ID = model_id
    return _EMBEDDER

def build_faiss_index(chunks: List[str], embed_model_id: str) -> Tuple[Any, Any]:
    if not chunks:
        raise ValueError("CV neobsahuje ziadny pouzitelny text/chunk.")

    embedder = get_embedder(embed_model_id)

    vectors = embedder.encode(
                                chunks,
                                convert_to_numpy=True,
                                normalize_embeddings=True,
                                show_progress_bar=True,
                            )

    np_mod = cast(Any, np)
    vectors = np_mod.asarray(vectors, dtype="float32")

    if vectors.ndim != 2:
        raise ValueError(f"Embedding model vratil necakany tvar vektorov: {vectors.shape}")

    faiss_mod = cast(Any, faiss)
    index = cast(Any, faiss_mod.IndexFlatIP(int(vectors.shape[1])))
    add_fn = cast(Any, index.add)
    add_fn(vectors)

    return index, vectors

def rag_search(query: str, chunks: List[str], index: Any, embed_model_id: str, top_k: int) -> List[str]:
    embedder = get_embedder(embed_model_id)
    q = embedder.encode([query], convert_to_numpy=True, normalize_embeddings=True, show_progress_bar=False).astype("float32")
    scores, ids = index.search(q, min(top_k, len(chunks)))
    results = []
    for score, idx in zip(scores[0], ids[0]):
        if idx < 0:
            continue
        if float(score) < MIN_RAG_SIMILARITY:
            continue
        results.append(f"[similarity={float(score):.3f}] {chunks[int(idx)]}")
    return results

# -----------------------------------------------------------------------------
# 6. REPORTING
# -----------------------------------------------------------------------------

def status_icon(status: str) -> str:
    return {"splnene": "✅","ciastocne_splnene": "🟡","nesplnene": "❌","nejasne": "⚪",}.get(status, "⚪")

def verdict(score: float) -> str:
    if score >= 80:
        return "Vhodny kandidat"
    if score >= 60:
        return "Skor vhodny kandidat"
    if score >= 40:
        return "Ciastocne vhodny / vyzaduje manualne posudenie"
    return "Slaba zhoda s poziciou"

def render_markdown_report(job_data: Dict[str, Any], candidate: Dict[str, Any], evals: List[Dict[str, Any]]) -> str:
    score = weighted_average(evals)
    lines = []
    lines.append("# AI CV validator")
    lines.append("")
    lines.append(f"**Pozicia:** {job_data.get('job_title', 'unknown')}")
    lines.append(f"**Seniorita:** {job_data.get('seniority', 'unknown')}")
    lines.append(f"**Zdroj poziadaviek:** {job_data.get('_source', 'unknown')}")
    lines.append(f"**Celkove skore:** {score:.2f} / 100")
    lines.append(f"**Odporucanie:** {verdict(score)}")
    lines.append("")
    lines.append("> Vystup je odporucanie pre cloveka, nie automaticke rozhodnutie o kandidatovi.")
    lines.append("")

    if candidate:
        lines.append("## Anonymizovany profil kandidata")
        lines.append(str(candidate.get("summary", "")))
        for key, title in [
            ("skills", "Zrucnosti"),
            ("experience", "Skusenosti"),
            ("education", "Vzdelanie"),
            ("languages", "Jazyky"),
            ("certifications", "Certifikaty"),
            ("risks_or_missing_info", "Rizika / chybajuce info"),
        ]:
            vals = candidate.get(key)
            if isinstance(vals, list) and vals:
                lines.append(f"\n**{title}:**")
                for v in vals[:12]:
                    lines.append(f"- {v}")
        lines.append("")

    lines.append("## Vyhodnotenie poziadaviek")
    lines.append("")
    lines.append("| Stav | Poziadavka | Skore | Priorita | Vysvetlenie |")
    lines.append("|---|---|---:|---|---|")
    for r in evals:
        exp = str(r.get("explanation", "")).replace("|", "/")
        req = str(r.get("requirement", "")).replace("|", "/")
        lines.append(
            f"| {status_icon(r.get('status', ''))} {r.get('status', '')} "
            f"| {req} | {float(r.get('score', 0)):.0f} "
            f"| {r.get('priority', '')} / w={r.get('weight', 1)} | {exp} |"
        )

    lines.append("")
    lines.append("## Odkazy a poznamky")
    for r in evals:
        lines.append(f"\n### {status_icon(r.get('status', ''))} {r.get('requirement_id', '')}: {r.get('requirement', '')}")
        if r.get("risk_note"):
            lines.append(f"**Riziko/neistota:** {r.get('risk_note')}")
        ev = r.get("evidence_used") or []
        if isinstance(ev, list) and ev:
            lines.append("**Pouzite odkazy:**")
            for e in ev[:3]:
                short = normalize_space(str(e))[:700]
                lines.append(f"- {short}")

    return "\n".join(lines)


# -----------------------------------------------------------------------------
# 6A. MLFLOW LOGGING HELPERS
# -----------------------------------------------------------------------------
_TOKENIZER_COUNTER_CACHE: Dict[str, Any] = {}


def _safe_mlflow_param(value: Any, max_len: int = 250) -> str:
    """Return a compact MLflow-safe string parameter value."""
    text = str(value if value is not None else "")
    text = normalize_space(text)
    if len(text) > max_len:
        return text[:max_len] + "..."
    return text


def _approx_token_count(text: str) -> int:
    """Very small fallback token estimate when tokenizer cannot be loaded."""
    text = str(text or "")
    if not text:
        return 0
    # Rough English/Slovak approximation: 1 token ~= 4 characters.
    return max(1, int(len(text) / 4))


def count_text_tokens(text: str, model_id: Optional[str] = None) -> int:
    """Count tokens using the selected model tokenizer, with safe fallback.

    The tokenizer is loaded only for counting and cached in memory. If it is not
    available locally, the function falls back to an approximate count so MLflow
    logging never breaks the validation pipeline.
    """
    text = str(text or "")
    if not text:
        return 0

    model_id = str(model_id or "").strip()
    if not model_id:
        return _approx_token_count(text)

    try:
        if model_id not in _TOKENIZER_COUNTER_CACHE:
            from transformers import AutoTokenizer

            _TOKENIZER_COUNTER_CACHE[model_id] = AutoTokenizer.from_pretrained(
                model_id,
                trust_remote_code=True,
                local_files_only=MLFLOW_TOKENIZER_LOCAL_ONLY,
            )

        tok = _TOKENIZER_COUNTER_CACHE[model_id]
        return len(tok(text, return_tensors=None, truncation=False)["input_ids"])

    except Exception:
        return _approx_token_count(text)


def _extract_metrics_from_json_report(json_report: str) -> Dict[str, float]:
    """Extract numeric metrics from the generated validation JSON."""
    metrics: Dict[str, float] = {}

    try:
        payload = json.loads(json_report or "{}")
    except Exception:
        return metrics

    overall_score = payload.get("overall_score")
    if isinstance(overall_score, (int, float)):
        metrics["overall_score"] = float(overall_score)

    evaluations = payload.get("evaluations", [])
    if isinstance(evaluations, list):
        metrics["requirement_count"] = float(len(evaluations))

        scores = []
        full_match = 0
        partial_match = 0
        missing = 0

        for item in evaluations:
            if not isinstance(item, dict):
                continue

            score = item.get("score")
            if isinstance(score, (int, float)):
                scores.append(float(score))

            status = str(item.get("status", "")).strip().lower()
            if "full" in status or "match" == status:
                full_match += 1
            elif "partial" in status:
                partial_match += 1
            elif "missing" in status or "no" in status:
                missing += 1

        if scores:
            metrics["average_requirement_score"] = float(sum(scores) / len(scores))
            metrics["min_requirement_score"] = float(min(scores))
            metrics["max_requirement_score"] = float(max(scores))

        metrics["full_match_count"] = float(full_match)
        metrics["partial_match_count"] = float(partial_match)
        metrics["missing_count"] = float(missing)

    return metrics


def log_validation_run_to_mlflow(
                                *,
                                cv_file: str,
                                job_url: str,
                                job_text_manual: str,
                                job_schema_xlsx_path: str,
                                manual_position_name: str,
                                model_id: str,
                                fallback_model_id: str,
                                aux_model_id: str,
                                load_mode: str,
                                embed_model_id: str,
                                top_k: int,
                                max_requirements: int,
                                include_candidate_summary: bool,
                                report_md: str,
                                json_report: str,
                                runtime_info: str,
                                elapsed_sec: float,
                                log_input_artifacts: bool = False,
                                ) -> str:
    """Log one CV validation run to MLflow.

    By default, sensitive source inputs such as full CV files and full job text are
    not logged. Enable `log_input_artifacts` only when you are allowed to store
    those artifacts in the local MLflow directory.
    """
    if not MLFLOW_READY or mlflow is None:
        raise RuntimeError(MLFLOW_SETUP_ERROR or "MLflow is not ready.")

    cv_file = str(cv_file or "")
    job_url = str(job_url or "")
    job_text_manual = str(job_text_manual or "")
    manual_position_name = str(manual_position_name or "")
    report_md = str(report_md or "")
    json_report = str(json_report or "")
    runtime_info = str(runtime_info or "")

    source_fingerprint = "|".join(
        [
            Path(cv_file).name if cv_file else "",
            job_url,
            manual_position_name,
            model_id,
            embed_model_id,
            str(top_k),
            str(max_requirements),
        ]
    )
    run_hash = hashlib.sha256(source_fingerprint.encode("utf-8", errors="ignore")).hexdigest()[:12]
    run_name = f"cv_validation_{run_hash}"

    input_text_for_count = "\n\n".join(
        [
            job_url,
            job_text_manual,
            manual_position_name,
            job_schema_xlsx_path,
        ]
    )
    output_text_for_count = "\n\n".join([report_md, json_report, runtime_info])

    input_tokens = count_text_tokens(input_text_for_count, model_id=model_id)
    output_tokens = count_text_tokens(output_text_for_count, model_id=model_id)

    metrics = {
        "input_tokens_estimated": float(input_tokens),
        "output_tokens_estimated": float(output_tokens),
        "total_tokens_estimated": float(input_tokens + output_tokens),
        "elapsed_sec": float(elapsed_sec),
        "job_text_chars": float(len(job_text_manual)),
        "report_chars": float(len(report_md)),
        "json_chars": float(len(json_report)),
    }
    metrics.update(_extract_metrics_from_json_report(json_report))

    with mlflow.start_run(run_name=run_name) as run:
        mlflow.log_params(
            {
                "app": "few-shot-cv-validator",
                "model_id": _safe_mlflow_param(model_id),
                "fallback_model_id": _safe_mlflow_param(fallback_model_id),
                "aux_model_id": _safe_mlflow_param(aux_model_id),
                "load_mode": _safe_mlflow_param(load_mode),
                "embed_model_id": _safe_mlflow_param(embed_model_id),
                "top_k": int(top_k),
                "max_requirements": int(max_requirements),
                "include_candidate_summary": bool(include_candidate_summary),
                "job_source": "manual_position" if manual_position_name else ("url" if job_url else "manual_text"),
                "manual_position_name": _safe_mlflow_param(manual_position_name),
                "job_url": _safe_mlflow_param(job_url),
                "cv_file_name": _safe_mlflow_param(Path(cv_file).name if cv_file else ""),
                "cv_file_ext": _safe_mlflow_param(file_ext(cv_file) if cv_file else ""),
                "log_input_artifacts": bool(log_input_artifacts),
                "run_hash": run_hash,
            }
        )

        mlflow.log_metrics(metrics)

        mlflow.set_tags(
            {
                "app": "few-shot-cv-validator",
                "task_type": "cv_job_matching",
                "prompt_mode": "few_shot_langchain",
            }
        )

        mlflow.log_text(report_md, "outputs/report.md")
        mlflow.log_text(json_report, "outputs/report.json")
        mlflow.log_text(runtime_info, "outputs/runtime_info.txt")

        run_metadata = {
            "model_id": model_id,
            "fallback_model_id": fallback_model_id,
            "aux_model_id": aux_model_id,
            "load_mode": load_mode,
            "embed_model_id": embed_model_id,
            "top_k": int(top_k),
            "max_requirements": int(max_requirements),
            "include_candidate_summary": bool(include_candidate_summary),
            "elapsed_sec": float(elapsed_sec),
            "input_tokens_estimated": input_tokens,
            "output_tokens_estimated": output_tokens,
            "mlflow_tracking_uri": mlflow.get_tracking_uri(),
            "mlflow_experiment_name": MLFLOW_EXPERIMENT_NAME,
        }
        mlflow.log_text(
            json.dumps(run_metadata, ensure_ascii=False, indent=2),
            "metadata/run_metadata.json",
        )

        if log_input_artifacts:
            if job_text_manual:
                mlflow.log_text(job_text_manual, "inputs/job_text_manual.txt")
            if job_url:
                mlflow.log_text(job_url, "inputs/job_url.txt")
            if manual_position_name:
                mlflow.log_text(manual_position_name, "inputs/manual_position_name.txt")
            if cv_file and Path(cv_file).exists():
                mlflow.log_artifact(cv_file, artifact_path="inputs/cv_file")

        return run.info.run_id


# -----------------------------------------------------------------------------
# 7. MAIN PIPELINE
# -----------------------------------------------------------------------------

def run_validation(
                    cv_file: str,
                    job_url: str,
                    job_text_manual: str,
                    job_schema_xlsx_path: str,
                    manual_position_name: str,
                    model_id: str,
                    fallback_model_id: str,
                    aux_model_id: str,
                    load_mode: str,
                    embed_model_id: str,
                    top_k: int,
                    max_requirements: int,
                    include_candidate_summary: bool,
) -> Tuple[str, str, str]:
    if not cv_file:
        raise gr.Error("Nahraj CV subor.")

    runtime = []
    runtime.append(cuda_summary())
    runtime.append(f"LLM model: {model_id}")
    runtime.append(f"Fallback model: {fallback_model_id}")
    runtime.append(f"Aux model: {aux_model_id or model_id}")
    runtime.append(f"Load mode: {load_mode}")
    runtime.append(f"Embedding model: {embed_model_id}")
    runtime.append(f"Top-K: {top_k}")
    runtime.append(f"Max requirements: {max_requirements}")
    runtime.append(f"Schema XLSX path: {job_schema_xlsx_path or DEFAULT_JOB_SCHEMA_XLSX_PATH}")
    runtime.append(f"Manual position: {manual_position_name or '-'}")
    runtime.append("Prompt mode: LangChain few-shot")

    cv_text = load_document(cv_file)
    if len(cv_text) < 100:
        raise gr.Error("Z CV sa podarilo vytiahnut velmi malo textu. Skus iny format, idealne PDF/DOCX.")
    runtime.append(f"CV text: {len(cv_text):,} znakov")

    job_requirement_schema_text, prompt_schema_source = load_job_requirement_schema_text(
        job_schema_xlsx_path or DEFAULT_JOB_SCHEMA_XLSX_PATH
    )
    runtime.append(f"Prompt schema source: {prompt_schema_source}")

    _, _, model_info = load_llm(model_id, load_mode, fallback_model_id)
    runtime.append(model_info)

    if manual_position_name and manual_position_name.strip():
        job_data = load_manual_job_requirements_from_excel(
                                                            position_query=manual_position_name,
                                                            schema_xlsx_path=job_schema_xlsx_path or DEFAULT_JOB_SCHEMA_XLSX_PATH,
                                                            max_requirements=max_requirements,
                                                            model_id=model_id,
                                                            load_mode=load_mode,
                                                            fallback_model_id=fallback_model_id,
                                                            aux_model_id=aux_model_id,
                                                            )
        runtime.append("Zdroj pozicie: manualny schema XLSX katalog")
    else:
        job_text = ""
        if job_text_manual and job_text_manual.strip():
            job_text = normalize_space(job_text_manual)
            runtime.append("Inzerat: pouzity manualne vlozeny text")
        elif job_url and job_url.strip():
            job_text = scrape_url(job_url)
            runtime.append(f"Inzerat: nacitany z URL, {len(job_text):,} znakov")
        else:
            raise gr.Error("Zadaj URL inzeratu, vloz text inzeratu manualne, alebo vypln manualnu poziciu zo schema XLSX.")

        if len(job_text) < 100:
            raise gr.Error("Z inzeratu sa podarilo ziskat velmi malo textu. Vloz text inzeratu manualne.")

        job_data = extract_job_requirements(
                                            job_text,
                                            job_requirement_schema_text,
                                            model_id,
                                            load_mode,
                                            fallback_model_id,
                                            aux_model_id,
                                            max_requirements,
                                            )
    requirements = job_data.get("requirements", [])
    if not requirements:
        job_preview = (job_text or "").strip()

        if len(job_preview) >= 100:
            requirements = [
                            {
                            "id": "REQ-001",
                            "requirement": "General match against the provided job advertisement",
                            "category": "general",
                            "priority": "medium",
                            "source": job_preview[:2000],
                            }
                            ]

            print(
                    "Warning: No structured requirements were extracted. "
                    "Using full job advertisement as one general fallback requirement."
                    )
        else:
            raise gr.Error(
                            "Inzerat je prazdny alebo prilis kratky. "
                            "Skontroluj, ci si vlozil spravny subor alebo cisty text inzeratu."
                            )
    runtime.append(f"Extrahovane poziadavky: {len(requirements)}")
    runtime.append(f"Zdroj poziadaviek: {job_data.get('_source', 'unknown')}")

    meta = job_data.get("_meta", {})
    if isinstance(meta, dict):
        runtime.append(f"LLM count: {meta.get('llm_count', 0)}")
        runtime.append(f"Fallback count: {meta.get('fallback_count', 0)}")
        runtime.append(f"Weak LLM: {meta.get('weak_llm', False)}")
        runtime.append(f"Merged count: {meta.get('merged_count', 0)}")
        runtime.append(f"Prompt mode meta: {meta.get('prompt_mode', 'few_shot_langchain')}")

    candidate = {}
    if include_candidate_summary:
        candidate = extract_candidate_summary(cv_text, model_id, load_mode, fallback_model_id)

    chunks = chunk_text(cv_text, CHUNK_WORDS, CHUNK_OVERLAP)
    index, _ = build_faiss_index(chunks, embed_model_id)
    runtime.append(f"CV chunks: {len(chunks)}")

    evals = []
    for req in requirements:
        evidence = rag_search(req.get("text", ""), chunks, index, embed_model_id, int(top_k))
        ev = evaluate_one_requirement(req, evidence, model_id, load_mode, fallback_model_id)
        evals.append(ev)

    final = {
                "job": job_data,
                "candidate_profile": candidate,
                "overall_score": weighted_average(evals),
                "verdict": verdict(weighted_average(evals)),
                "evaluations": evals,
                "runtime": runtime,
            }

    md = render_markdown_report(job_data, candidate, evals)
    js = json.dumps(final, ensure_ascii=False, indent=2)
    return md, js, "\n".join(runtime)

def gradio_run_wrapper(*args):
    """Gradio wrapper with optional MLflow tracking.

    The last two UI inputs are:
    - mlflow_enabled
    - mlflow_log_input_artifacts

    They are intentionally handled here so the core validation pipeline stays
    focused on validation logic.
    """
    try:
        if len(args) < 15:
            return run_validation(*args)

        validation_args = args[:13]
        mlflow_enabled = bool(args[13])
        mlflow_log_input_artifacts = bool(args[14])

        started_at = time.perf_counter()
        report_md, json_report, runtime_info = run_validation(*validation_args)
        elapsed_sec = time.perf_counter() - started_at

        runtime_info = f"{runtime_info}\nElapsed: {elapsed_sec:.2f} sec"

        if mlflow_enabled and MLFLOW_READY:
            if not MLFLOW_READY:
                runtime_info = (
                    f"{runtime_info}\n"
                    f"MLflow logging skipped: {MLFLOW_SETUP_ERROR or 'MLflow is not ready.'}"
                )
                return report_md, json_report, runtime_info

            try:
                run_id = log_validation_run_to_mlflow(
                    cv_file=validation_args[0],
                    job_url=validation_args[1],
                    job_text_manual=validation_args[2],
                    job_schema_xlsx_path=validation_args[3],
                    manual_position_name=validation_args[4],
                    model_id=validation_args[5],
                    fallback_model_id=validation_args[6],
                    aux_model_id=validation_args[7],
                    load_mode=validation_args[8],
                    embed_model_id=validation_args[9],
                    top_k=int(validation_args[10]),
                    max_requirements=int(validation_args[11]),
                    include_candidate_summary=bool(validation_args[12]),
                    report_md=report_md,
                    json_report=json_report,
                    runtime_info=runtime_info,
                    elapsed_sec=elapsed_sec,
                    log_input_artifacts=mlflow_log_input_artifacts,
                )
                runtime_info = f"{runtime_info}\nMLflow run logged: {run_id}"

            except Exception as log_exc:
                runtime_info = (
                    f"{runtime_info}\n"
                    f"MLflow logging failed: {type(log_exc).__name__}: {log_exc}"
                )

        return report_md, json_report, runtime_info

    except gr.Error:
        raise
    except Exception as exc:
        err = f"Chyba: {type(exc).__name__}: {exc}\n\n{traceback.format_exc()}"
        return "# Chyba pri spracovani\n\n```text\n" + err + "\n```", "{}", err

# -----------------------------------------------------------------------------
# 8. GRADIO UI
# -----------------------------------------------------------------------------

def build_ui():
    with gr.Blocks(title="Lokalny AI CV Validator") as demo:
        gr.Markdown("")

        with gr.Row():
            with gr.Column(scale=1):
                cv_file = gr.File(
                    label="CV subor",
                    file_types=[".pdf", ".docx", ".doc", ".rtf", ".txt", ".md"],
                    type="filepath",
                )
                job_url = gr.Textbox(label="URL inzeratu", placeholder="https://")
                job_text_manual = gr.Textbox(
                    label="Alebo vloz text inzeratu manualne",
                    lines=8,
                    placeholder="Text pracovnej ponuky",
                )

                with gr.Accordion("Externy schema XLSX", open=False):
                    job_schema_xlsx_path = gr.Textbox(
                        label="Schema XLSX path",
                        value=DEFAULT_JOB_SCHEMA_XLSX_PATH,
                    )
                    manual_position_name = gr.Textbox(
                        label="Manualna pozicia zo schema XLSX",
                        placeholder="napr. python_backend_medior alebo Python developer",
                    )

                with gr.Accordion("Model nastavenia", open=False):
                    model_id = gr.Textbox(label="LLM model z Hugging Face", value=DEFAULT_LLM_MODEL_ID)
                    fallback_model_id = gr.Textbox(label="Fallback LLM model", value=DEFAULT_FALLBACK_LLM_MODEL_ID)
                    aux_model_id = gr.Textbox(label="Aux LLM model pre canonicalizaciu/genericnost", value=DEFAULT_AUX_LLM_MODEL_ID or DEFAULT_LLM_MODEL_ID)
                    load_mode = gr.Dropdown(
                        label="Load mode",
                        choices=["auto", "bnb_4bit", "fp16_gpu", "cpu"],
                        value=LLM_LOAD_MODE,
                    )
                    embed_model_id = gr.Textbox(label="Embedding model z Hugging Face", value=DEFAULT_EMBED_MODEL_ID)
                    top_k = gr.Slider(label="Top-K dokazov z CV", minimum=2, maximum=20, step=1, value=DEFAULT_TOP_K)
                    max_requirements = gr.Slider(
                        label="Max pocet poziadaviek z inzeratu",
                        minimum=5,
                        maximum=25,
                        step=1,
                        value=DEFAULT_MAX_REQUIREMENTS,
                    )
                    include_candidate_summary = gr.Checkbox(label="Extrahovat anonymizovany profil kandidata", value=True)

                with gr.Accordion("MLflow experiment tracking", open=False):
                    mlflow_status_text = (
                        "MLflow ready - local SQLite tracking is enabled."
                        if MLFLOW_READY
                        else f"MLflow not ready - {MLFLOW_SETUP_ERROR or 'install mlflow first'}"
                    )
                    gr.Markdown(
                        f"""
                        **Prompt / validation benchmarking**

                        MLflow can store each validation run as an experiment:
                        model settings, generated report, JSON output, estimated token usage and runtime.

                        Status: `{mlflow_status_text}`

                        Local UI command:
                        `mlflow ui --backend-store-uri sqlite:///{MLFLOW_DB_PATH.as_posix()}`

                        Optional auto-start from script:
                        set `START_MLFLOW_UI=1` and optionally `OPEN_MLFLOW_UI_BROWSER=1`.
                        """
                    )

                    mlflow_enabled = gr.Checkbox(
                        label="Enable MLflow logging",
                        value=False,
                        interactive=MLFLOW_READY,
                    )
                    mlflow_log_input_artifacts = gr.Checkbox(
                        label="Also log source inputs/CV file as artifacts - use only for non-sensitive data",
                        value=False,
                        interactive=MLFLOW_READY,
                    )

                with gr.Row():
                    run_btn = gr.Button("Spustit validaciu", variant="primary")
                    unload_btn = gr.Button("Uvolnit model z VRAM")

            with gr.Column(scale=2):
                report_md = gr.Markdown(label="Report")
                runtime_info = gr.Textbox(label="Runtime info", lines=8)
                json_report = gr.Code(label="JSON report", language="json", lines=20)

        run_btn.click(
            fn=gradio_run_wrapper,
            inputs=[
                    cv_file,
                    job_url,
                    job_text_manual,
                    job_schema_xlsx_path,
                    manual_position_name,
                    model_id,
                    fallback_model_id,
                    aux_model_id,
                    load_mode,
                    embed_model_id,
                    top_k,
                    max_requirements,
                    include_candidate_summary,
                    mlflow_enabled,
                    mlflow_log_input_artifacts,
                    ],
            outputs=[report_md, json_report, runtime_info],
        )
        unload_btn.click(fn=unload_llm, inputs=[], outputs=[runtime_info])

    return demo

if __name__ == "__main__":
    demo = build_ui()
    demo.launch(server_name="127.0.0.1", server_port=7860, inbrowser=True)