# app-few_shot.py of nyx3ton-project\CV Validator

from __future__ import annotations

# -----------------------------------------------------------------------------
# 0. IMPORTS
# 1. ENV + GLOBAL SETTINGS
# 2. GLOBAL MODEL CACHE
# 3. UTILS
# 4. DOCUMENT LOADING: PDF/DOCX/RTF/TXT/DOC
# 5. JOB AD SCRAPING
# 6. CHUNKING + RAG
# 7. LOCAL HUGGING FACE LLM
# 8. FEW-SHOT EXAMPLES + LANGCHAIN PROMPTS
# 9. LLM PROMPTS / TASKS
# 10. MAIN PIPELINE
# 11. GRADIO UI
# -----------------------------------------------------------------------------

# -----------------------------------------------------------------------------
# 0. IMPORTS
# -----------------------------------------------------------------------------
import gc, hashlib, json, os, re, tempfile, traceback, unicodedata, torch, requests, faiss
from dataclasses import dataclass
from importlib import import_module
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, cast
import numpy as np
import gradio as gr
from bs4 import BeautifulSoup
from sentence_transformers import SentenceTransformer
from transformers import AutoModelForCausalLM, AutoTokenizer, BitsAndBytesConfig
from dotenv import load_dotenv
from openpyxl import load_workbook

from langchain_core.prompts import ChatPromptTemplate, FewShotChatMessagePromptTemplate

from dictionary_fallback import (fallback_extract_requirements_from_text, build_hybrid_requirement_result)

# -----------------------------------------------------------------------------
# 1. ENV + GLOBAL SETTINGS
# -----------------------------------------------------------------------------

APP_DIR = Path(__file__).resolve().parent
load_dotenv(APP_DIR / ".env")

def env_bool(name: str, default: bool = False) -> bool:
    value = os.getenv(name)
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "y", "on"}

DEFAULT_LLM_MODEL_ID = os.getenv("LLM_MODEL_ID", "unsloth/Qwen3.5-4B")
DEFAULT_FALLBACK_LLM_MODEL_ID = os.getenv("FALLBACK_LLM_MODEL_ID", "Qwen/Qwen2.5-3B-Instruct")
DEFAULT_AUX_LLM_MODEL_ID = os.getenv("AUX_LLM_MODEL_ID", "").strip()
DEFAULT_EMBED_MODEL_ID = os.getenv("EMBED_MODEL_ID","sentence-transformers/paraphrase-multilingual-mpnet-base-v2")
DEFAULT_JOB_SCHEMA_XLSX_PATH = os.getenv("JOB_SCHEMA_XLSX_PATH", str(APP_DIR / "job_requirement_schema.xlsx"))
DEFAULT_JOB_SCHEMA_SHEET = os.getenv("JOB_SCHEMA_SHEET", "ManualPositions")
DEFAULT_PROMPT_SCHEMA_SHEET = os.getenv("PROMPT_SCHEMA_SHEET", "PromptSchema")

LLM_LOAD_MODE = os.getenv("LLM_LOAD_MODE", "fp16_gpu")  # auto | bnb_4bit | fp16_gpu | cpu
MAX_GPU_MEMORY = os.getenv("MAX_GPU_MEMORY", "10.5GiB")
MAX_INPUT_TOKENS = int(os.getenv("MAX_INPUT_TOKENS", "8192"))
MAX_NEW_TOKENS = int(os.getenv("MAX_NEW_TOKENS", "900"))
CHUNK_WORDS = int(os.getenv("CHUNK_WORDS", "220"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "55"))
DEFAULT_TOP_K = max(2, min(20, int(os.getenv("TOP_K", "5"))))
DEFAULT_MAX_REQUIREMENTS = max(5, min(25, int(os.getenv("MAX_REQUIREMENTS", "12"))))
MIN_RAG_SIMILARITY = float(os.getenv("MIN_RAG_SIMILARITY", "0.20"))
TEMPERATURE_SETTING = float(os.getenv("DEF_TEMPERATURE_SETTING", "0.2"))
SAMPLE_SETTING = env_bool("DEF_SAMPLE_SETTING", False)
P_SETTING = float(os.getenv("DEF_P_SETTING", "0.95"))
GEN_TOP_K_SETTING = int(os.getenv("DEF_TOP_K_SETTING", "20"))
REPETITION_PEN = float(os.getenv("DEF_REPETITION_PEN", "1.0"))

ENABLE_LLM_GENERICNESS = env_bool("ENABLE_LLM_GENERICNESS", True)
ENABLE_LLM_CANONICALIZATION = env_bool("ENABLE_LLM_CANONICALIZATION", True)
GENERIC_HEURISTIC_LOW = float(os.getenv("GENERIC_HEURISTIC_LOW", "0.25"))
GENERIC_HEURISTIC_HIGH = float(os.getenv("GENERIC_HEURISTIC_HIGH", "0.75"))
MIN_LLM_CANONICAL_TEXT_LEN = int(os.getenv("MIN_LLM_CANONICAL_TEXT_LEN", "18"))

HF_HOME_LOCAL = os.getenv("HF_HOME_LOCAL", "").strip()
if HF_HOME_LOCAL:
    os.environ["HF_HOME"] = str(Path(HF_HOME_LOCAL).expanduser().resolve())

# -----------------------------------------------------------------------------
# 2. GLOBAL MODEL CACHE
# -----------------------------------------------------------------------------

_TOKENIZER = None
_MODEL = None
_MODEL_INFO = "Model este nie je nacitany."
_EMBEDDER = None
_EMBEDDER_ID = None
_GENERICNESS_CACHE: Dict[str, Dict[str, Any]] = {}
_CANONICAL_CACHE: Dict[str, Dict[str, Any]] = {}

@dataclass
class Requirement:
    id: str
    text: str
    category: str = "other"
    priority: str = "unknown"
    weight: float = 1.0

# -----------------------------------------------------------------------------
# 3. UTILS
# -----------------------------------------------------------------------------

def strip_thinking(text: str) -> str:
    if not text:
        return text

    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL | re.IGNORECASE).strip()

    if "</think>" in text:
        text = text.split("</think>", 1)[-1].strip()

    return text.strip()

def safe_json_loads(text: str, fallback: Any) -> Any:
    if not text:
        return fallback

    cleaned = strip_thinking(text).strip()
    cleaned = re.sub(r"^```(?:json)?", "", cleaned, flags=re.IGNORECASE).strip()
    cleaned = re.sub(r"```$", "", cleaned).strip()

    try:
        return json.loads(cleaned)
    except Exception:
        pass

    decoder = json.JSONDecoder()
    parsed_candidates = []

    for i, ch in enumerate(cleaned):
        if ch not in "[{":
            continue
        try:
            obj, _ = decoder.raw_decode(cleaned[i:])
            parsed_candidates.append(obj)
        except Exception:
            continue

    if parsed_candidates:
        return parsed_candidates[-1]

    return fallback

def normalize_space(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()

def trim_text(text: str, max_chars: int = 24000) -> str:
    text = text.strip()
    if len(text) <= max_chars:
        return text
    half = max_chars // 2
    return text[:half] + "\n\n[...SKRATENE...]\n\n" + text[-half:]

def file_ext(path: str) -> str:
    return Path(path).suffix.lower().replace(".", "")

def weighted_average(rows: List[Dict[str, Any]]) -> float:
    total_w = 0.0
    total = 0.0
    for r in rows:
        try:
            w = float(r.get("weight", 1.0))
            s = float(r.get("score", 0.0))
        except Exception:
            continue
        total_w += max(w, 0.01)
        total += max(0.0, min(100.0, s)) * max(w, 0.01)
    return round(total / total_w, 2) if total_w else 0.0

def remove_diacritics(text: str) -> str:
    normalized = unicodedata.normalize("NFKD", text)
    return "".join(ch for ch in normalized if not unicodedata.combining(ch))

def cache_key(prefix: str, text: str, model_id: str = "") -> str:
    raw = f"{prefix}||{model_id}||{normalize_space(text)}"
    return hashlib.sha1(raw.encode("utf-8", errors="ignore")).hexdigest()

def parse_boolish(value: Any, default: bool = True) -> bool:
    if value is None:
        return default
    if isinstance(value, bool):
        return value
    text = normalize_space(str(value)).lower()
    if not text:
        return default
    return text in {"1", "true", "yes", "y", "on", "ano"}

def normalized_lookup_value(value: Any) -> str:
    return remove_diacritics(normalize_space(str(value or "")).lower())

def safe_float_value(value: Any, default: float = 1.0) -> float:
    try:
        return float(value)
    except Exception:
        return default

def load_job_requirement_schema_text(schema_xlsx_path: str) -> Tuple[str, str]:
    path = Path(schema_xlsx_path or "").expanduser()

    if not schema_xlsx_path or not path.exists():
        return DEFAULT_JOB_REQUIREMENT_SCHEMA, "builtin_fallback"

    try:
        wb = load_workbook(filename=str(path), read_only=True, data_only=True)
        if DEFAULT_PROMPT_SCHEMA_SHEET not in wb.sheetnames:
            return DEFAULT_JOB_REQUIREMENT_SCHEMA, "builtin_fallback_missing_prompt_sheet"

        ws = wb[DEFAULT_PROMPT_SCHEMA_SHEET]
        schema_text = ""

        for row in ws.iter_rows(min_row=2, values_only=True):
            key = normalized_lookup_value(row[0] if len(row) > 0 else "")
            value = row[1] if len(row) > 1 else ""
            if key in {"schema_json", "job_requirement_schema", "json_schema"} and str(value or "").strip():
                schema_text = str(value).strip()
                break

        if schema_text:
            return schema_text, f"excel:{path}"

    except Exception:
        pass

    return DEFAULT_JOB_REQUIREMENT_SCHEMA, "builtin_fallback_error"

def load_manual_job_requirements_from_excel(
                                            position_query: str,
                                            schema_xlsx_path: str,
                                            max_requirements: int,
                                            model_id: str,
                                            load_mode: str,
                                            fallback_model_id: str,
                                            aux_model_id: str,
) -> Dict[str, Any]:
    path = Path(schema_xlsx_path or "").expanduser()

    if not position_query or not position_query.strip():
        raise gr.Error("Pre manualnu poziciu zo schema XLSX zadaj nazov pozicie alebo position_key.")

    if not schema_xlsx_path or not path.exists():
        raise gr.Error("Schema XLSX subor neexistuje. Skontroluj JOB_SCHEMA_XLSX_PATH alebo pole v UI.")

    try:
        wb = load_workbook(filename=str(path), read_only=True, data_only=True)
    except Exception as exc:
        raise gr.Error(f"Nepodarilo sa otvorit schema XLSX: {exc}") from exc

    if DEFAULT_JOB_SCHEMA_SHEET not in wb.sheetnames:
        raise gr.Error(f"V schema XLSX chyba harok '{DEFAULT_JOB_SCHEMA_SHEET}'.")

    ws = wb[DEFAULT_JOB_SCHEMA_SHEET]
    rows_iter = ws.iter_rows(min_row=1, max_row=1, values_only=True)
    headers_row = next(rows_iter, None)
    if not headers_row:
        raise gr.Error("Harok ManualPositions nema hlavicku.")

    normalized_headers = [normalized_lookup_value(h) for h in headers_row]

    header_aliases = {
                        "position_key": {"position_key", "position id", "position_id", "role_key", "job_key"},
                        "job_title": {"job_title", "job title", "pozicia", "position_name", "role_name", "nazov_pozicie"},
                        "seniority": {"seniority", "seniorita", "level"},
                        "requirement_id": {"requirement_id", "requirement id", "req_id", "id"},
                        "requirement_text": {"requirement_text", "requirement text", "text", "poziadavka", "requirement"},
                        "category": {"category", "kategoria"},
                        "priority": {"priority", "priorita"},
                        "weight": {"weight", "vaha"},
                        "active": {"active", "aktivne", "enabled", "is_active"},
                        "note": {"note", "notes", "poznamka"},
                        }

    header_positions: Dict[str, int] = {}
    for canonical_name, variants in header_aliases.items():
        for idx, header_name in enumerate(normalized_headers):
            if header_name in variants:
                header_positions[canonical_name] = idx
                break

    required_headers = {"job_title", "requirement_text"}
    missing = [x for x in required_headers if x not in header_positions]
    if missing:
        raise gr.Error(
            "Harok ManualPositions nema povinne stlpce: " + ", ".join(missing) + "."
        )

    grouped_rows: Dict[str, List[Dict[str, Any]]] = {}

    for row in ws.iter_rows(min_row=2, values_only=True):
        if row is None:
            continue

        def get_value(name: str, default: Any = "") -> Any:
            idx = header_positions.get(name)
            if idx is None or idx >= len(row):
                return default
            return row[idx]

        requirement_text = normalize_space(str(get_value("requirement_text", "")))
        if not requirement_text:
            continue

        active = parse_boolish(get_value("active", True), default=True)
        if not active:
            continue

        position_key = normalize_space(str(get_value("position_key", "")))
        job_title = normalize_space(str(get_value("job_title", "")))
        seniority = normalize_space(str(get_value("seniority", "unknown"))) or "unknown"

        logical_group = position_key or job_title
        if not logical_group:
            continue

        grouped_rows.setdefault(logical_group, []).append(
                                                            {
                                                            "position_key": position_key,
                                                            "job_title": job_title or logical_group,
                                                            "seniority": seniority,
                                                            "requirement_id": normalize_space(str(get_value("requirement_id", ""))),
                                                            "requirement_text": requirement_text,
                                                            "category": normalize_space(str(get_value("category", "other"))) or "other",
                                                            "priority": normalize_space(str(get_value("priority", "unknown"))) or "unknown",
                                                            "weight": safe_float_value(get_value("weight", 1.0), 1.0),
                                                            "note": normalize_space(str(get_value("note", ""))),
                                                            }
                                                        )

    if not grouped_rows:
        raise gr.Error("Schema XLSX neobsahuje ziadne aktivne manualne poziadavky.")

    query_norm = normalized_lookup_value(position_query)
    best_group_name = ""
    best_score = -1

    for group_name, group_rows in grouped_rows.items():
        sample = group_rows[0]
        title_norm = normalized_lookup_value(sample.get("job_title", ""))
        key_norm = normalized_lookup_value(sample.get("position_key", ""))
        group_norm = normalized_lookup_value(group_name)

        score = 0
        if query_norm and query_norm == key_norm:
            score = 100
        elif query_norm and query_norm == title_norm:
            score = 98
        elif query_norm and query_norm == group_norm:
            score = 96
        elif query_norm and key_norm and query_norm in key_norm:
            score = 90
        elif query_norm and title_norm and query_norm in title_norm:
            score = 88
        elif query_norm and group_norm and query_norm in group_norm:
            score = 86
        elif query_norm and key_norm and key_norm in query_norm:
            score = 80
        elif query_norm and title_norm and title_norm in query_norm:
            score = 78

        if score > best_score:
            best_group_name = group_name
            best_score = score

    if best_score < 0 or best_group_name not in grouped_rows:
        raise gr.Error(f"V schema XLSX sa nenasla manualna pozicia pre '{position_query}'.")

    matched_rows = grouped_rows[best_group_name]
    sample = matched_rows[0]

    job_data = {
                "job_title": sample.get("job_title") or best_group_name,
                "seniority": sample.get("seniority") or "unknown",
                "requirements": [],
                "_source": "external_excel_schema",
                "_meta": {
                            "schema_path": str(path.resolve()),
                            "schema_sheet": DEFAULT_JOB_SCHEMA_SHEET,
                            "manual_position_query": position_query,
                            "matched_group": best_group_name,
                            "match_score": best_score,
                            "prompt_mode": "manual_excel_schema",
                        },
                }

    raw_requirements = []
    for i, item in enumerate(matched_rows, start=1):
        raw_requirements.append(
            {
                "id": item.get("requirement_id") or f"R{i}",
                "text": item.get("requirement_text", ""),
                "category": item.get("category", "other"),
                "priority": item.get("priority", "unknown"),
                "weight": item.get("weight", 1.0),
                "note": item.get("note", ""),
            }
        )

    job_data["requirements"] = clean_requirements(
                                                    raw_requirements,
                                                    max_requirements,
                                                    model_id=model_id,
                                                    load_mode=load_mode,
                                                    fallback_model_id=fallback_model_id,
                                                    aux_model_id=aux_model_id,
                                                    allow_llm=True,
                                                )

    return job_data

def normalize_requirement_key_basic(text: str) -> str:
    text = remove_diacritics(normalize_space(text).lower())
    text = text.replace("&", " and ")
    text = re.sub(r"[^a-zA-Z0-9+.#/ ]+", " ", text)

    replacements = {
                    "programovanie v jazyku python": "python",
                    "skusenosti s pythonom": "python",
                    "pokrocila znalost python": "python",
                    "znalost python": "python",
                    "programovanie v jazyku java": "java",
                    "skusenosti s javou": "java",
                    "znalost sql": "sql",
                    "skusenosti so sql": "sql",
                    "anglicky jazyk": "anglictina",
                    "anglictina b2": "anglictina na urovni b2",
                    "english b2": "anglictina na urovni b2",
                    "timova praca": "teamova spolupraca",
                    "praca v time": "teamova spolupraca",
                    "ms office": "microsoft office",
                    "ms excel": "microsoft excel",
                    "microsoft excel": "microsoft excel",
                    "excel": "microsoft excel",
                    "aj": "anglictina",
                    "nj": "nemcina",
                    }

    stop_phrases = [
                    "znalost",
                    "prakticka znalost",
                    "pokrocila znalost",
                    "vyborna znalost",
                    "zakladna znalost",
                    "skusenosti s",
                    "skusenost s",
                    "prax s",
                    "prax v",
                    "schopnost",
                    "schopnost pracovat s",
                    "ovladanie",
                    "aktivna",
                    "vyhodou je",
                    "podmienkou je",
                    "skusenosti v oblasti",
                    ]

    text = replacements.get(text, text)

    for phrase in stop_phrases:
        text = re.sub(rf"\b{re.escape(phrase)}\b", " ", text)

    text = re.sub(r"\s+", " ", text).strip()

    direct_aliases = {
                        "python developer": "python",
                        "python vyvoj": "python",
                        "rest apis": "rest api",
                        "api": "api",
                        "sql databazy": "sql",
                        "english language": "anglictina",
                        "nemecky jazyk": "nemcina",
                        }

    text = direct_aliases.get(text, text)
    return text

def direct_canonical_requirement(text: str) -> Optional[Dict[str, Any]]:
    key = normalize_requirement_key_basic(text)

    mapping = {
                "python": ("python", "Python", "hard_skill"),
                "java": ("java", "Java", "hard_skill"),
                "sql": ("sql", "SQL", "hard_skill"),
                "rest api": ("rest api", "REST API", "hard_skill"),
                "docker": ("docker", "Docker", "hard_skill"),
                "kubernetes": ("kubernetes", "Kubernetes", "hard_skill"),
                "git": ("git", "Git", "hard_skill"),
                "linux": ("linux", "Linux", "hard_skill"),
                "aws": ("aws", "AWS", "hard_skill"),
                "azure": ("azure", "Azure", "hard_skill"),
                "sap successfactors": ("sap successfactors", "SAP SuccessFactors", "hard_skill"),
                "microsoft excel": ("microsoft excel", "Microsoft Excel", "hard_skill"),
                "microsoft office": ("microsoft office", "Microsoft Office", "hard_skill"),
                "anglictina": ("anglictina", "anglictina", "language"),
                "anglictina na urovni b2": ("anglictina na urovni b2", "anglictina na urovni B2", "language"),
                "nemcina": ("nemcina", "nemcina", "language"),
                "teamova spolupraca": ("teamova spolupraca", "teamova spolupraca", "soft_skill"),
                "komunikacne schopnosti": ("komunikacne schopnosti", "komunikacne schopnosti", "soft_skill"),
                "samostatnost": ("samostatnost", "samostatnost", "soft_skill"),
                "zodpovednost": ("zodpovednost", "zodpovednost", "soft_skill"),
                "bratislava": ("bratislava", "Bratislava", "location"),
                }

    if key in mapping:
        canonical_key, canonical_text, category_hint = mapping[key]
        return {
                "canonical_key": canonical_key,
                "canonical_text": canonical_text,
                "category_hint": category_hint,
                "confidence": 0.99,
                "reason": "Priama heuristicka zhoda.",
                "source": "heuristic_direct",
                }

    return None

def heuristic_genericness_score(text: str) -> Dict[str, Any]:
    original = normalize_space(text)
    key = normalize_requirement_key_basic(original)
    words = key.split()

    generic_patterns = [
                        "komunikacne schopnosti",
                        "komunikacia",
                        "samostatnost",
                        "zodpovednost",
                        "flexibilita",
                        "odolnost voci stresu",
                        "teamova spolupraca",
                        "proaktivny pristup",
                        "motivacia",
                        "ucit sa nove veci",
                        "prijemne vystupovanie",
                        "spolahlivost",
                        ]

    concrete_patterns = [
                        "python",
                        "java",
                        "sql",
                        "excel",
                        "microsoft excel",
                        "docker",
                        "kubernetes",
                        "linux",
                        "aws",
                        "azure",
                        "sap",
                        "rest api",
                        "api",
                        "git",
                        "b2",
                        "c1",
                        "c2",
                        "bratislava",
                        "praxe",
                        "rok",
                        "roky",
                        "certifikat",
                        "vzdelanie",
                        "bakalar",
                        "magister",
                        "doktorand"
                        ]

    score = 0.50
    reasons = []

    if any(p == key or p in key for p in generic_patterns):
        score += 0.32
        reasons.append("Obsahuje vseobecnu soft-skill formulaciu.")

    if len(words) <= 2 and not re.search(r"\d", key):
        score += 0.12
        reasons.append("Poziadavka je velmi kratka a bez meratelneho detailu.")

    if re.search(r"\b(aspon|min\.?|minimalne)\b", key):
        score -= 0.18
        reasons.append("Obsahuje minimalnu alebo meratelnu podmienku.")

    if re.search(r"\b(a1|a2|b1|b2|c1|c2)\b", key):
        score -= 0.24
        reasons.append("Obsahuje explicitnu jazykovu uroven.")

    if re.search(r"\b\d+\b", key):
        score -= 0.20
        reasons.append("Obsahuje cislo alebo kvantifikaciu.")

    if any(p in key for p in concrete_patterns):
        score -= 0.30
        reasons.append("Obsahuje konkretny nastroj, technologiu alebo meratelny signal.")

    if len(words) >= 6:
        score -= 0.08
        reasons.append("Poziadavka je opisnejsia a konkretnejsia.")

    score = max(0.0, min(1.0, score))
    confidence = 0.55 + abs(score - 0.5)
    confidence = max(0.0, min(0.99, confidence))

    return {
            "is_generic": score >= 0.50,
            "confidence": round(confidence, 4),
            "heuristic_score": round(score, 4),
            "reason": " ".join(reasons) if reasons else "Bez silneho heuristickeho signalu.",
            "source": "heuristic",
            }

def resolve_assist_model_id(primary_model_id: str, aux_model_id: Optional[str] = None) -> str:
    return (aux_model_id or DEFAULT_AUX_LLM_MODEL_ID or primary_model_id).strip()

def normalize_requirement_key(
                                text: str,
                                model_id: Optional[str] = None,
                                load_mode: Optional[str] = None,
                                fallback_model_id: Optional[str] = None,
                                aux_model_id: Optional[str] = None,
                                allow_llm: bool = True,
) -> str:
    data = hybrid_normalize_requirement_key(
                                            text=text,
                                            model_id=model_id or DEFAULT_LLM_MODEL_ID,
                                            load_mode=load_mode or LLM_LOAD_MODE,
                                            fallback_model_id=fallback_model_id or DEFAULT_FALLBACK_LLM_MODEL_ID,
                                            aux_model_id=aux_model_id,
                                            allow_llm=allow_llm,
                                            )
    return str(data.get("canonical_key") or normalize_requirement_key_basic(text))

def dedupe_text_list(items: List[str], limit: int = 12) -> List[str]:
    seen = set()
    result = []
    for item in items:
        item = normalize_space(str(item))
        if not item:
            continue
        key = item.lower()
        if key in seen:
            continue
        seen.add(key)
        result.append(item)
        if len(result) >= limit:
            break
    return result

def llm_classify_requirement_genericness(
                                        text: str,
                                        model_id: str,
                                        load_mode: str,
                                        fallback_model_id: str,
                                        aux_model_id: Optional[str] = None,
) -> Dict[str, Any]:
    assist_model_id = resolve_assist_model_id(model_id, aux_model_id)
    ck = cache_key("genericness", text, assist_model_id)
    if ck in _GENERICNESS_CACHE:
        return _GENERICNESS_CACHE[ck]

    final_user = (
                    "Posud, ci je poziadavka na kandidata prilis vseobecna alebo naopak dost konkretna.\n\n"
                    "Vrat presne tento JSON tvar:\n"
                    "{\n"
                    '  "is_generic": true,\n'
                    '  "confidence": 0.0,\n'
                    '  "reason": "kratke vysvetlenie",\n'
                    '  "suggested_focus": "ak je to vhodne, kratke konkretne jadro poziadavky, inak prazdny string"\n'
                    "}\n\n"
                    "Pravidla:\n"
                    "- true pouzi pri vseobecnych soft-skill formulaciach bez konkretneho obsahu\n"
                    "- false pouzi pri technologii, jazyku, meratelnej poziadavke, nastroji, praxi alebo lokalite\n"
                    "- vrat iba JSON\n\n"
                    "POZIADAVKA:\n"
                    + normalize_space(text)
                    )

    messages = build_lc_messages_from_fewshot(system_text=SYSTEM_REQUIREMENT_UTILS_JSON,examples=GENERICNESS_FEWSHOT,example_human_template="Posud generickost poziadavky.\n\nPOZIADAVKA:\n{input}",final_human_input=final_user)
    raw = chat_generate_messages(messages,assist_model_id,load_mode,fallback_model_id,max_new_tokens=220,do_sample=False)
    data = safe_json_loads(raw,
        fallback={
            "is_generic": False,
            "confidence": 0.0,
            "reason": "LLM klasifikacia zlyhala.",
            "suggested_focus": "",
        },
    )

    if not isinstance(data, dict):
        data = {
            "is_generic": False,
            "confidence": 0.0,
            "reason": "LLM klasifikacia zlyhala.",
            "suggested_focus": "",
        }

    result = {
                "is_generic": bool(data.get("is_generic", False)),
                "confidence": max(0.0, min(1.0, float(data.get("confidence", 0.0) or 0.0))),
                "heuristic_score": None,
                "reason": normalize_space(str(data.get("reason") or "")),
                "suggested_focus": normalize_space(str(data.get("suggested_focus") or "")),
                "source": "llm_genericness",
            }

    _GENERICNESS_CACHE[ck] = result
    return result

def score_requirement_genericness(
                                    text: str,
                                    model_id: Optional[str] = None,
                                    load_mode: Optional[str] = None,
                                    fallback_model_id: Optional[str] = None,
                                    aux_model_id: Optional[str] = None,
                                    allow_llm: bool = True,
) -> Dict[str, Any]:
    heuristic = heuristic_genericness_score(text)

    if heuristic["heuristic_score"] <= GENERIC_HEURISTIC_LOW:
        return {
            **heuristic,
            "is_generic": False,
            "source": "heuristic_strong_non_generic",
        }

    if heuristic["heuristic_score"] >= GENERIC_HEURISTIC_HIGH:
        return {
                **heuristic,"is_generic": True,"source": "heuristic_strong_generic",
                }

    if not allow_llm or not ENABLE_LLM_GENERICNESS:
        return heuristic

    return llm_classify_requirement_genericness(text=text,model_id=model_id or DEFAULT_LLM_MODEL_ID,load_mode=load_mode or LLM_LOAD_MODE,fallback_model_id=fallback_model_id or DEFAULT_FALLBACK_LLM_MODEL_ID,aux_model_id=aux_model_id)

def is_generic_requirement_text(
                                text: str,
                                model_id: Optional[str] = None,
                                load_mode: Optional[str] = None,
                                fallback_model_id: Optional[str] = None,
                                aux_model_id: Optional[str] = None,
                                allow_llm: bool = True,
) -> bool:
    data = score_requirement_genericness(
                                        text=text,
                                        model_id=model_id,
                                        load_mode=load_mode,
                                        fallback_model_id=fallback_model_id,
                                        aux_model_id=aux_model_id,
                                        allow_llm=allow_llm,
                                        )
    return bool(data.get("is_generic", False))

def llm_canonicalize_requirement(
                                text: str,
                                model_id: str,
                                load_mode: str,
                                fallback_model_id: str,
                                aux_model_id: Optional[str] = None,
) -> Dict[str, Any]:
    assist_model_id = resolve_assist_model_id(model_id, aux_model_id)
    ck = cache_key("canonical", text, assist_model_id)
    if ck in _CANONICAL_CACHE:
        return _CANONICAL_CACHE[ck]

    final_user = (
                    "Znormalizuj jednu poziadavku na kandidata do kanonickeho kluca vhodneho na deduplikaciu.\n\n"
                    "Vrat presne tento JSON tvar:\n"
                    "{\n"
                    '  "canonical_key": "kratky kluc malymi pismenami",\n'
                    '  "canonical_text": "kratka citatelna verzia",\n'
                    '  "category_hint": "hard_skill|experience|education|language|soft_skill|location|other|unknown",\n'
                    '  "confidence": 0.0,\n'
                    '  "reason": "kratke vysvetlenie"\n'
                    "}\n\n"
                    "Pravidla:\n"
                    "- canonical_key musi byt kratky, stabilny a bez zbytocnych slov\n"
                    "- zachovaj hlavny skill, jazyk, lokalitu alebo jadro poziadavky\n"
                    "- nepis vetu, ale jadro poziadavky\n"
                    "- vrat iba JSON\n\n"
                    "POZIADAVKA:\n"
                    + normalize_space(text)
                )

    messages = build_lc_messages_from_fewshot(system_text=SYSTEM_REQUIREMENT_UTILS_JSON,examples=CANONICALIZATION_FEWSHOT,example_human_template="Znormalizuj poziadavku do kanonickej formy.\n\nPOZIADAVKA:\n{input}",final_human_input=final_user)

    raw = chat_generate_messages(messages,assist_model_id,load_mode,fallback_model_id,max_new_tokens=260,do_sample=False)

    data = safe_json_loads(raw,
                            fallback={
                                        "canonical_key": normalize_requirement_key_basic(text),
                                        "canonical_text": normalize_space(text),
                                        "category_hint": "unknown",
                                        "confidence": 0.0,"reason": "LLM kanonizacia zlyhala."},
                            )

    if not isinstance(data, dict):
        data = {
                "canonical_key": normalize_requirement_key_basic(text),
                "canonical_text": normalize_space(text),
                "category_hint": "unknown",
                "confidence": 0.0,
                "reason": "LLM kanonizacia zlyhala.",
                }

    canonical_key = normalize_requirement_key_basic(str(data.get("canonical_key") or text))
    canonical_text = normalize_space(str(data.get("canonical_text") or text))
    category_hint = str(data.get("category_hint") or "unknown")
    if category_hint not in {"hard_skill", "experience", "education", "language", "soft_skill", "location", "other", "unknown"}:
        category_hint = "unknown"

    result = {
                "canonical_key": canonical_key or normalize_requirement_key_basic(text),
                "canonical_text": canonical_text or normalize_space(text),
                "category_hint": category_hint,
                "confidence": max(0.0, min(1.0, float(data.get("confidence", 0.0) or 0.0))),
                "reason": normalize_space(str(data.get("reason") or "")),
                "source": "llm_canonicalizer",
                }

    _CANONICAL_CACHE[ck] = result
    return result

def hybrid_normalize_requirement_key(
                                    text: str,
                                    model_id: Optional[str] = None,
                                    load_mode: Optional[str] = None,
                                    fallback_model_id: Optional[str] = None,
                                    aux_model_id: Optional[str] = None,
                                    allow_llm: bool = True,
) -> Dict[str, Any]:
    text = normalize_space(text)
    direct = direct_canonical_requirement(text)
    if direct is not None:
        return direct

    basic_key = normalize_requirement_key_basic(text)
    heuristic_result = {
                        "canonical_key": basic_key,
                        "canonical_text": text,
                        "category_hint": "unknown",
                        "confidence": 0.68 if len(basic_key.split()) <= 3 else 0.52,
                        "reason": "Heuristicka kanonizacia bez LLM.",
                        "source": "heuristic",
                        }

    if not allow_llm or not ENABLE_LLM_CANONICALIZATION:
        return heuristic_result

    if len(text) < MIN_LLM_CANONICAL_TEXT_LEN and len(basic_key.split()) <= 3:
        return heuristic_result

    llm_result = llm_canonicalize_requirement(text=text,
                                                model_id=model_id or DEFAULT_LLM_MODEL_ID,
                                                load_mode=load_mode or LLM_LOAD_MODE,
                                                fallback_model_id=fallback_model_id or DEFAULT_FALLBACK_LLM_MODEL_ID,
                                                aux_model_id=aux_model_id,
                                                )

    if not llm_result.get("canonical_key"):
        return heuristic_result

    return llm_result

def clean_requirements(
                        reqs: List[Any],
                        max_requirements: int,
                        model_id: Optional[str] = None,
                        load_mode: Optional[str] = None,
                        fallback_model_id: Optional[str] = None,
                        aux_model_id: Optional[str] = None,
                        allow_llm: bool = True,
) -> List[Dict[str, Any]]:
    clean_reqs = []
    seen_keys = set()

    for i, r in enumerate(reqs[: max_requirements * 3], start=1):
        if not isinstance(r, dict):
            continue

        text = normalize_space(str(r.get("text", "")))
        if not text:
            continue

        canonical = hybrid_normalize_requirement_key(
                                                    text=text,
                                                    model_id=model_id,
                                                    load_mode=load_mode,
                                                    fallback_model_id=fallback_model_id,
                                                    aux_model_id=aux_model_id,
                                                    allow_llm=allow_llm,
                                                    )

        key = str(canonical.get("canonical_key") or normalize_requirement_key_basic(text))
        if not key or key in seen_keys:
            continue
        seen_keys.add(key)

        try:
            weight = float(r.get("weight", 1.0))
        except Exception:
            weight = 1.0

        category = str(r.get("category") or "other")
        if category not in {"hard_skill", "experience", "education", "language", "soft_skill", "location", "other"}:
            category = "other"

        category_hint = str(canonical.get("category_hint") or "unknown")
        if category == "other" and category_hint in {"hard_skill", "experience", "education", "language", "soft_skill", "location"}:
            category = category_hint

        priority = str(r.get("priority") or "unknown")
        if priority not in {"must", "nice", "unknown"}:
            priority = "unknown"

        generic_meta = score_requirement_genericness(
                                                    text=text,
                                                    model_id=model_id,
                                                    load_mode=load_mode,
                                                    fallback_model_id=fallback_model_id,
                                                    aux_model_id=aux_model_id,
                                                    allow_llm=allow_llm,
                                                    )

        clean_reqs.append(
            {
                "id": str(r.get("id") or f"R{i}"),
                "text": text,
                "category": category,
                "priority": priority,
                "weight": max(0.5, min(5.0, weight)),
                "canonical_key": key,
                "generic_flag": bool(generic_meta.get("is_generic", False)),
            }
        )

        if len(clean_reqs) >= max_requirements:
            break

    for idx, item in enumerate(clean_reqs, start=1):
        item["id"] = f"R{idx}"

    return clean_reqs

def is_weak_requirement_output(
                                data: Dict[str, Any],
                                max_requirements: int,
                                model_id: Optional[str] = None,
                                load_mode: Optional[str] = None,
                                fallback_model_id: Optional[str] = None,
                                aux_model_id: Optional[str] = None,
                                allow_llm: bool = True,
) -> bool:
    reqs = data.get("requirements") or []
    if not isinstance(reqs, list):
        return True

    if len(reqs) < min(3, max_requirements):
        return True

    generic_count = 0
    hard_like_count = 0
    seen = set()
    dup_count = 0

    for req in reqs:
        if not isinstance(req, dict):
            continue

        text = normalize_space(str(req.get("text", "")))
        key = str(req.get("canonical_key") or normalize_requirement_key(
            text,
            model_id=model_id,
            load_mode=load_mode,
            fallback_model_id=fallback_model_id,
            aux_model_id=aux_model_id,
            allow_llm=allow_llm,
        ))

        if not key:
            continue

        if key in seen:
            dup_count += 1
        seen.add(key)

        if bool(req.get("generic_flag")) or is_generic_requirement_text(
                                                                        text,
                                                                        model_id=model_id,
                                                                        load_mode=load_mode,
                                                                        fallback_model_id=fallback_model_id,
                                                                        aux_model_id=aux_model_id,
                                                                        allow_llm=allow_llm,
                                                                        ):
            generic_count += 1

        category = str(req.get("category") or "other")
        if category in {"hard_skill", "experience", "education", "language"}:
            hard_like_count += 1

    if dup_count >= 2:
        return True
    if generic_count >= max(2, len(reqs) // 2):
        return True
    if hard_like_count == 0:
        return True

    return False

# -----------------------------------------------------------------------------
# 4. DOCUMENT LOADING: PDF/DOCX/RTF/TXT/DOC
# -----------------------------------------------------------------------------

def load_pdf(path: str) -> str:
    fitz = import_module("fitz")  # PyMuPDF

    parts = []
    with fitz.open(path) as doc:
        page_count = int(getattr(doc, "page_count", 0))

        for page_index in range(page_count):
            page = doc.load_page(page_index)
            text = page.get_text("text") or ""

            if text.strip():
                parts.append(f"\n--- PAGE {page_index + 1} ---\n{text}")

    return normalize_space("\n".join(parts))

def load_docx(path: str) -> str:
    import docx

    d = docx.Document(path)
    paragraphs = [p.text for p in d.paragraphs if p.text and p.text.strip()]

    for table in d.tables:
        for row in table.rows:
            cells = [normalize_space(c.text) for c in row.cells if c.text]
            if cells:
                paragraphs.append(" | ".join(cells))

    return normalize_space("\n".join(paragraphs))

def load_rtf(path: str) -> str:
    from striprtf.striprtf import rtf_to_text

    raw = Path(path).read_text(errors="ignore")
    return normalize_space(rtf_to_text(raw))

def load_txt(path: str) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp1250", "latin-1"):
        try:
            return normalize_space(Path(path).read_text(encoding=enc, errors="ignore"))
        except Exception:
            continue
    return normalize_space(Path(path).read_text(errors="ignore"))

def load_doc_legacy_windows(path: str) -> str:
    try:
        import win32com.client  # type: ignore
    except Exception as exc:
        raise RuntimeError(
            "Subor .doc je legacy format. Pre native Windows fallback treba mat nainstalovany "
            "Microsoft Word + pywin32. Alternativa: uloz CV ako .docx alebo .pdf."
        ) from exc

    tmp_dir = Path(tempfile.mkdtemp(prefix="cvdoc_"))
    tmp_txt = tmp_dir / "converted.txt"
    word = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(Path(path).resolve()))
        doc.SaveAs(str(tmp_txt), FileFormat=7)  # 7 = wdFormatUnicodeText
        doc.Close(False)
        return load_txt(str(tmp_txt))
    finally:
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass

def load_document(path: str) -> str:
    ext = file_ext(path)
    if ext == "pdf":
        return load_pdf(path)
    if ext == "docx":
        return load_docx(path)
    if ext == "rtf":
        return load_rtf(path)
    if ext in {"txt", "md"}:
        return load_txt(path)
    if ext == "doc":
        return load_doc_legacy_windows(path)
    raise ValueError(f"Nepodporovany format suboru: .{ext}. Pouzi PDF, DOCX, RTF, TXT alebo DOC.")

# -----------------------------------------------------------------------------
# 5. JOB AD SCRAPING
# -----------------------------------------------------------------------------

def scrape_url(url: str) -> str:
    if not url or not url.strip():
        return ""

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122 Safari/537.36"
        )
    }
    resp = requests.get(url.strip(), headers=headers, timeout=25)
    resp.raise_for_status()

    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup(["script", "style", "noscript", "svg", "nav", "footer", "header"]):
        tag.decompose()

    main = soup.find("main") or soup.find("article") or soup.body or soup
    text = main.get_text("\n")
    lines = [normalize_space(x) for x in text.splitlines()]
    lines = [x for x in lines if len(x) > 2]

    seen = set()
    unique = []
    for line in lines:
        key = line.lower()
        if key in seen:
            continue
        seen.add(key)
        unique.append(line)

    return normalize_space("\n".join(unique))

# -----------------------------------------------------------------------------
# 6. CHUNKING + RAG
# -----------------------------------------------------------------------------

def chunk_text(text: str, words_per_chunk: int = CHUNK_WORDS, overlap: int = CHUNK_OVERLAP) -> List[str]:
    words = text.split()
    if not words:
        return []
    chunks = []
    step = max(1, words_per_chunk - overlap)
    for start in range(0, len(words), step):
        chunk_words = words[start:start + words_per_chunk]
        if len(chunk_words) < 20 and chunks:
            break
        chunks.append(" ".join(chunk_words))
    return chunks

def get_embedder(model_id: str = DEFAULT_EMBED_MODEL_ID):
    global _EMBEDDER, _EMBEDDER_ID
    if _EMBEDDER is not None and _EMBEDDER_ID == model_id:
        return _EMBEDDER
    device = "cpu"
    _EMBEDDER = SentenceTransformer(model_id, device=device)
    _EMBEDDER_ID = model_id
    return _EMBEDDER

def build_faiss_index(chunks: List[str], embed_model_id: str) -> Tuple[Any, Any]:
    if not chunks:
        raise ValueError("CV neobsahuje ziadny pouzitelny text/chunk.")

    embedder = get_embedder(embed_model_id)

    vectors = embedder.encode(
                                chunks,
                                convert_to_numpy=True,
                                normalize_embeddings=True,
                                show_progress_bar=True,
                            )

    np_mod = cast(Any, np)
    vectors = np_mod.asarray(vectors, dtype="float32")

    if vectors.ndim != 2:
        raise ValueError(f"Embedding model vratil necakany tvar vektorov: {vectors.shape}")

    faiss_mod = cast(Any, faiss)
    index = cast(Any, faiss_mod.IndexFlatIP(int(vectors.shape[1])))
    add_fn = cast(Any, index.add)
    add_fn(vectors)

    return index, vectors

def rag_search(query: str, chunks: List[str], index: Any, embed_model_id: str, top_k: int) -> List[str]:
    embedder = get_embedder(embed_model_id)
    q = embedder.encode([query], convert_to_numpy=True, normalize_embeddings=True, show_progress_bar=False).astype("float32")
    scores, ids = index.search(q, min(top_k, len(chunks)))
    results = []
    for score, idx in zip(scores[0], ids[0]):
        if idx < 0:
            continue
        if float(score) < MIN_RAG_SIMILARITY:
            continue
        results.append(f"[similarity={float(score):.3f}] {chunks[int(idx)]}")
    return results

# -----------------------------------------------------------------------------
# 7. LOCAL HUGGING FACE LLM
# -----------------------------------------------------------------------------

def cuda_summary() -> str:
    if not torch.cuda.is_available():
        return "CUDA nie je dostupna, pouzije sa CPU alebo fallback."
    name = torch.cuda.get_device_name(0)
    total_gb = torch.cuda.get_device_properties(0).total_memory / 1024 ** 3
    allocated_gb = torch.cuda.memory_allocated(0) / 1024 ** 3
    reserved_gb = torch.cuda.memory_reserved(0) / 1024 ** 3
    return f"CUDA OK: {name}, VRAM total={total_gb:.1f} GB, allocated={allocated_gb:.2f} GB, reserved={reserved_gb:.2f} GB"

def unload_llm() -> str:
    global _TOKENIZER, _MODEL, _MODEL_INFO
    _TOKENIZER = None
    _MODEL = None
    _MODEL_INFO = "Model bol uvolneny."
    gc.collect()
    if torch.cuda.is_available():
        torch.cuda.empty_cache()
    return _MODEL_INFO + "\n" + cuda_summary()

def load_llm(model_id: str = DEFAULT_LLM_MODEL_ID,load_mode: str = LLM_LOAD_MODE,fallback_model_id: str = DEFAULT_FALLBACK_LLM_MODEL_ID):
    global _TOKENIZER, _MODEL, _MODEL_INFO

    desired_signature = f"{model_id}|{load_mode}|fallback={fallback_model_id}"
    if _MODEL is not None and _TOKENIZER is not None and desired_signature in _MODEL_INFO:
        return _TOKENIZER, _MODEL, _MODEL_INFO

    unload_llm()

    errors = []
    has_cuda = torch.cuda.is_available()

    def _load_tokenizer(mid: str):
        return AutoTokenizer.from_pretrained(mid, trust_remote_code=True)

    def _try_4bit(mid: str):
        if not has_cuda:
            raise RuntimeError("CUDA nie je dostupna pre 4-bit GPU load.")
        bnb_config = BitsAndBytesConfig(
                                        load_in_4bit=True,
                                        bnb_4bit_quant_type="nf4",
                                        bnb_4bit_compute_dtype=torch.float16,
                                        bnb_4bit_use_double_quant=True,
                                        )
        tok = _load_tokenizer(mid)
        mdl = AutoModelForCausalLM.from_pretrained(
                                                    mid,
                                                    quantization_config=bnb_config,
                                                    device_map="auto",
                                                    max_memory={0: MAX_GPU_MEMORY, "cpu": "48GiB"},
                                                    trust_remote_code=True,
                                                    )
        return tok, mdl, f"Nacitany model: {mid} | mode=bnb_4bit | {desired_signature}"

    def _try_fp16_gpu(mid: str):
        if not has_cuda:
            raise RuntimeError("CUDA nie je dostupna pre fp16_gpu load.")
        tok = _load_tokenizer(mid)
        mdl = AutoModelForCausalLM.from_pretrained(
                                                    mid,
                                                    torch_dtype=torch.float16,
                                                    device_map="auto",
                                                    max_memory={0: MAX_GPU_MEMORY, "cpu": "48GiB"},
                                                    trust_remote_code=True,
                                                    )
        return tok, mdl, f"Nacitany model: {mid} | mode=fp16_gpu | {desired_signature}"

    def _try_cpu(mid: str):
        tok = _load_tokenizer(mid)
        mdl = AutoModelForCausalLM.from_pretrained(
                                                    mid,
                                                    torch_dtype=torch.float32,
                                                    device_map={"": "cpu"},
                                                    trust_remote_code=True,
                                                    )
        return tok, mdl, f"Nacitany model: {mid} | mode=cpu | {desired_signature}"

    attempts = []
    mode = (load_mode or "auto").lower().strip()

    if mode == "bnb_4bit":
        attempts = [lambda: _try_4bit(model_id)]
    elif mode == "fp16_gpu":
        attempts = [lambda: _try_fp16_gpu(model_id)]
    elif mode == "cpu":
        attempts = [lambda: _try_cpu(model_id)]
    else:
        attempts = [
                    lambda: _try_4bit(model_id),
                    lambda: _try_fp16_gpu(model_id),
                    lambda: _try_fp16_gpu(fallback_model_id),
                    lambda: _try_cpu(fallback_model_id),
                    ]

    for attempt in attempts:
        try:
            _TOKENIZER, _MODEL, _MODEL_INFO = attempt()
            _MODEL.eval()
            return _TOKENIZER, _MODEL, _MODEL_INFO + "\n" + cuda_summary()
        except Exception as exc:
            errors.append(f"{type(exc).__name__}: {exc}")
            gc.collect()
            if torch.cuda.is_available():
                torch.cuda.empty_cache()

    raise RuntimeError(
        "Nepodarilo sa nacitat lokalny LLM.\n\nPokusy:\n- " + "\n- ".join(errors)
    )

def model_device(model: Any):
    try:
        return next(model.parameters()).device
    except Exception:
        return "cuda" if torch.cuda.is_available() else "cpu"

def lc_messages_to_hf_messages(messages: List[Any]) -> List[Dict[str, str]]:
    role_map = {
                "system": "system",
                "human": "user",
                "ai": "assistant",
                }

    converted = []
    for msg in messages:
        msg_type = getattr(msg, "type", "human")
        role = role_map.get(msg_type, "user")
        content = normalize_space(str(getattr(msg, "content", "")))
        if content:
            converted.append({"role": role, "content": content})
    return converted

def chat_generate_messages(
                            messages: List[Dict[str, str]],
                            model_id: str,
                            load_mode: str,
                            fallback_model_id: str,
                            max_new_tokens: int = MAX_NEW_TOKENS,
                            do_sample: bool = False,
                            temperature: float = 0.2,
) -> str:
    tok, mdl, _ = load_llm(model_id, load_mode, fallback_model_id)

    if getattr(tok, "chat_template", None):
        prompt = tok.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)
    else:
        parts = []
        for m in messages:
            parts.append(f"{m['role'].capitalize()}:\n{m['content']}")
        prompt = "\n\n".join(parts) + "\n\nAssistant:\n"

    inputs = tok(prompt, return_tensors="pt", truncation=True, max_length=MAX_INPUT_TOKENS)
    dev = model_device(mdl)
    inputs = {k: v.to(dev) for k, v in inputs.items()}

    generation_kwargs = {
                        **inputs,
                        "max_new_tokens": max_new_tokens,
                        "pad_token_id": tok.eos_token_id,
                        "repetition_penalty": REPETITION_PEN,
                        "do_sample": do_sample,
                        }

    if do_sample:
                generation_kwargs["temperature"] = temperature
                generation_kwargs["top_p"] = P_SETTING
                generation_kwargs["top_k"] = GEN_TOP_K_SETTING

    with torch.no_grad():
        out = mdl.generate(**generation_kwargs)

    generated = out[0][inputs["input_ids"].shape[-1]:]
    text = tok.decode(generated, skip_special_tokens=True)
    return strip_thinking(text)

SYSTEM_JSON = """
                Si lokalny AI asistent pre validaciu zivotopisov voci pracovnemu inzeratu.

                Tvoje pravidla:
                - odpovedaj iba validnym JSON objektom alebo validnym JSON polom
                - nepouzivaj markdown
                - nepouzivaj text mimo JSON
                - nehadaj informacie, ktore nie su v dodanom texte
                - pri hodnoteni kandidata pouzivaj iba dodane odkazy z CV
                - nevykonavaj finalne rozhodnutie o prijati kandidata
                - vystup sluzi iba ako odporucanie pre cloveka
                - nehodnot citlive atributy ako vek, pohlavie, narodnost, zdravotny stav, rodinny stav, fotografia alebo adresa
                - pis slovensky bez diakritiky
                """.strip()

SYSTEM_REQUIREMENT_UTILS_JSON = """
                Si lokalny AI asistent pre mikro-ulohy pri spracovani poziadaviek z pracovnych inzeratov.

                Tvoje pravidla:
                - odpovedaj iba validnym JSON objektom
                - nepouzivaj markdown
                - nepouzivaj text mimo JSON
                - bud konzistentny a kratky
                - canonical_key pis malymi pismenami bez zbytocnych slov
                - pis slovensky bez diakritiky
                """.strip()

# -----------------------------------------------------------------------------
# 8. FEW-SHOT EXAMPLES + LANGCHAIN PROMPTS
# -----------------------------------------------------------------------------
DEFAULT_JOB_REQUIREMENT_SCHEMA = """
{
    "job_title": "",
    "seniority": "junior|medior|senior|unknown",
    "requirements": [
        {
            "id": "R1",
            "text": "konkretna poziadavka",
            "category": "hard_skill|experience|education|language|soft_skill|location|other",
            "priority": "must|nice|unknown",
            "weight": 1.0
        }
    ]
}
""".strip()

CANDIDATE_SUMMARY_SCHEMA = """
{
    "summary": "kratke zhrnutie kandidata bez mena a citlivych atributov",
    "skills": ["..."],
    "experience": ["..."],
    "education": ["..."],
    "languages": ["..."],
    "certifications": ["..."],
    "risks_or_missing_info": ["..."]
}
""".strip()

REQUIREMENT_EVAL_SCHEMA = {
    "requirement_id": "R1",
    "requirement": "konkretna poziadavka",
    "status": "splnene|ciastocne_splnene|nesplnene|nejasne",
    "score": 0,
    "confidence": 0.0,
    "evidence_used": ["kratke citacie/parafrazy dokazov"],
    "explanation": "kratke vysvetlenie v slovencine bez diakritiky",
    "risk_note": "co chyba alebo preco je hodnotenie neiste"
}

JOB_REQ_FEWSHOT = [
    {
        "input": """Hladame Python developera pre backendovy tim. Pozadujeme Python, SQL, Git a REST API. Nutna je prax aspon 2 roky s vyvojom aplikacii. Vyhodou je Docker a Kubernetes. Ocenime anglictinu na urovni B2. Ponukame flexibilny home office, multisport kartu a teambuildingy.""",
        "output": {
            "job_title": "Python developer",
            "seniority": "medior",
            "requirements": [
                {"id": "R1", "text": "Python", "category": "hard_skill", "priority": "must", "weight": 5.0},
                {"id": "R2", "text": "SQL", "category": "hard_skill", "priority": "must", "weight": 4.0},
                {"id": "R3", "text": "Git", "category": "hard_skill", "priority": "must", "weight": 3.0},
                {"id": "R4", "text": "REST API", "category": "hard_skill", "priority": "must", "weight": 3.0},
                {"id": "R5", "text": "aspon 2 roky praxe s vyvojom aplikacii", "category": "experience", "priority": "must", "weight": 4.0},
                {"id": "R6", "text": "Docker", "category": "hard_skill", "priority": "nice", "weight": 2.0},
                {"id": "R7", "text": "Kubernetes", "category": "hard_skill", "priority": "nice", "weight": 2.0},
                {"id": "R8", "text": "anglictina na urovni B2", "category": "language", "priority": "nice", "weight": 2.0}
            ]
        },
    },
    {
        "input": """Pre klienta hladame HR specialistu. Podmienkou je skusenost s naborom kandidatok a kandidatov, aktivna anglictina, dobra znalost Excelu a komunikacne schopnosti. Vyhodou je skusenost so SAP SuccessFactors. Pozicia je v Bratislave. Firemne benefity a kultura nie su sucast poziadaviek.""",
        "output": {
            "job_title": "HR specialista",
            "seniority": "unknown",
            "requirements": [
                {"id": "R1", "text": "skusenost s naborom kandidatok a kandidatov", "category": "experience", "priority": "must", "weight": 4.0},
                {"id": "R2", "text": "aktivna anglictina", "category": "language", "priority": "must", "weight": 3.0},
                {"id": "R3", "text": "Excel", "category": "hard_skill", "priority": "must", "weight": 3.0},
                {"id": "R4", "text": "komunikacne schopnosti", "category": "soft_skill", "priority": "must", "weight": 3.0},
                {"id": "R5", "text": "SAP SuccessFactors", "category": "hard_skill", "priority": "nice", "weight": 2.0},
                {"id": "R6", "text": "Bratislava", "category": "location", "priority": "unknown", "weight": 1.0}
            ]
        },
    },
]

CANDIDATE_FEWSHOT = [
    {
        "input": """Jan Novak
Backend developer
5 rokov vyvoja v Pythone a Django. Skusenosti s SQL, REST API, Dockerom a Gitom. Anglictina B2. Vzdelanie: STU FEI. Certifikat AWS Cloud Practitioner.""",
        "output": {
            "summary": "Kandidat ma viacrocnu prax v backendovom vyvoji so zameranim na Python technologie.",
            "skills": ["Python", "Django", "SQL", "REST API", "Docker", "Git"],
            "experience": ["5 rokov backendovy vyvoj"],
            "education": ["STU FEI"],
            "languages": ["anglictina B2"],
            "certifications": ["AWS Cloud Practitioner"],
            "risks_or_missing_info": ["Nie je jasna uroven prace s Kubernetes."]
        },
    }
]

REQUIREMENT_EVAL_FEWSHOT = [
    {
        "input": {
            "requirement": {"id": "R1", "text": "Python", "category": "hard_skill", "priority": "must", "weight": 5.0},
            "evidence": [
                "[similarity=0.882] 5 rokov vyvoja v Pythone a Django.",
                "[similarity=0.771] Backend developer so zameranim na Python aplikacie."
            ]
        },
        "output": {
            "requirement_id": "R1",
            "requirement": "Python",
            "status": "splnene",
            "score": 96,
            "confidence": 0.95,
            "evidence_used": [
                "5 rokov vyvoja v Pythone a Django.",
                "Backend developer so zameranim na Python aplikacie."
            ],
            "explanation": "CV obsahuje priamy a silny dokaz o viacrocnej praxi s Pythonom.",
            "risk_note": ""
        },
    },
    {
        "input": {
            "requirement": {"id": "R2", "text": "anglictina na urovni B2", "category": "language", "priority": "must", "weight": 3.0},
            "evidence": [
                "[similarity=0.691] Active use of English in international environment.",
                "[similarity=0.655] Daily communication with foreign clients."
            ]
        },
        "output": {
            "requirement_id": "R2",
            "requirement": "anglictina na urovni B2",
            "status": "ciastocne_splnene",
            "score": 72,
            "confidence": 0.74,
            "evidence_used": [
                "Active use of English in international environment.",
                "Daily communication with foreign clients."
            ],
            "explanation": "CV naznacuje prakticke pouzivanie anglictiny, ale priamo neuvadza uroven B2.",
            "risk_note": "Chyba explicitne uvedena jazykova uroven."
        },
    },
    {
        "input": {
            "requirement": {"id": "R3", "text": "Kubernetes", "category": "hard_skill", "priority": "nice", "weight": 2.0},
            "evidence": []
        },
        "output": {
            "requirement_id": "R3",
            "requirement": "Kubernetes",
            "status": "nesplnene",
            "score": 10,
            "confidence": 0.92,
            "evidence_used": [],
            "explanation": "V dodanych odkazoch z CV nie je dokaz o znalosti Kubernetes.",
            "risk_note": "Poziadavka nemala ziadny relevantny dokaz v CV."
        },
    },
]

GENERICNESS_FEWSHOT = [
    {
        "input": "komunikacne schopnosti",
        "output": {
            "is_generic": True,
            "confidence": 0.97,
            "reason": "Ide o vseobecnu soft-skill formulaciu bez konkretneho kontextu.",
            "suggested_focus": "komunikacia s klientom"
        },
    },
    {
        "input": "anglictina na urovni B2",
        "output": {
            "is_generic": False,
            "confidence": 0.98,
            "reason": "Poziadavka obsahuje konkretnu jazykovu uroven.",
            "suggested_focus": ""
        },
    },
    {
        "input": "aspon 2 roky praxe s navrhom SQL dotazov",
        "output": {
            "is_generic": False,
            "confidence": 0.99,
            "reason": "Obsahuje konkretnu technologiu aj meratelnu prax.",
            "suggested_focus": ""
        },
    },
]

CANONICALIZATION_FEWSHOT = [
    {
        "input": "prakticke skusenosti s programovanim v jazyku Python",
        "output": {
            "canonical_key": "python",
            "canonical_text": "Python",
            "category_hint": "hard_skill",
            "confidence": 0.98,
            "reason": "Hlavnym jadrom poziadavky je znalost jazyka Python."
        },
    },
    {
        "input": "aktivna anglictina na komunikativnej urovni B2",
        "output": {
            "canonical_key": "anglictina na urovni b2",
            "canonical_text": "anglictina na urovni B2",
            "category_hint": "language",
            "confidence": 0.99,
            "reason": "Poziadavka je jazykova a explicitne obsahuje uroven B2."
        },
    },
    {
        "input": "schopnost analyzovat business poziadavky a prekladat ich do technickych zadani",
        "output": {
            "canonical_key": "analyza business poziadaviek",
            "canonical_text": "analyza business poziadaviek",
            "category_hint": "other",
            "confidence": 0.87,
            "reason": "Jadrom je analyza a transformacia business poziadaviek do technickej formy."
        },
    },
]

def build_lc_messages_from_fewshot(
                                    system_text: str,
                                    examples: List[Dict[str, Any]],
                                    example_human_template: str,
                                    final_human_input: str,
) -> List[Dict[str, str]]:
    example_prompt = ChatPromptTemplate.from_messages([("human", example_human_template),("ai", "{output}")])

    normalized_examples = []
    for ex in examples:
        normalized_examples.append({**ex,"output": json.dumps(ex["output"], ensure_ascii=False, indent=2)})

    few_shot_prompt = FewShotChatMessagePromptTemplate(
                                                        example_prompt=example_prompt,
                                                        examples=normalized_examples,
                                                        )

    final_prompt = ChatPromptTemplate.from_messages([("system", system_text),few_shot_prompt,("human", final_human_input)])
    prompt_value = final_prompt.format_prompt()
    return lc_messages_to_hf_messages(prompt_value.to_messages())

# -----------------------------------------------------------------------------
# 9. LLM PROMPTS / TASKS
# -----------------------------------------------------------------------------

def extract_job_requirements(
                            job_text: str,
                            job_requirement_schema_text: str,
                            model_id: str,
                            load_mode: str,
                            fallback_model_id: str,
                            aux_model_id: str,
                            max_requirements: int,
) -> Dict[str, Any]:
    final_user = (
                    "Z pracovnej ponuky extrahuj atomicke poziadavky na kandidata.\n\n"
                    "Vrat presne tento JSON tvar:\n"
                    + job_requirement_schema_text
                    + "\n\nPravidla:\n"
                    f"- maximalne {max_requirements} najdolezitejsich poziadaviek\n"
                    "- must-have poziadavky daj weight 3 az 5\n"
                    "- nice-to-have poziadavky daj weight 1 az 2\n"
                    "- nerozbijaj jednu poziadavku na duplicitne varianty\n"
                    "- ignoruj benefity firmy, marketing a pravne formulacie\n"
                    "- ak nevies urcit job_title alebo seniority, pouzi hodnotu unknown\n"
                    "- vrat iba JSON, bez komentara\n\n"
                    "PRACOVNA PONUKA:\n"
                    + trim_text(job_text, 22000)
                    )

    messages = build_lc_messages_from_fewshot(
                                                system_text=SYSTEM_JSON,
                                                examples=JOB_REQ_FEWSHOT,
                                                example_human_template="Z pracovnej ponuky extrahuj atomicke poziadavky na kandidata.\n\nPRACOVNA PONUKA:\n{input}",
                                                final_human_input=final_user,
                                            )

    raw = chat_generate_messages(
                                messages,
                                model_id,
                                load_mode,
                                fallback_model_id,
                                max_new_tokens=1100,
                                do_sample=False,
                                )

    print("\n===== RAW JOB REQUIREMENTS OUTPUT =====")
    print(raw)
    print("===== END RAW JOB REQUIREMENTS OUTPUT =====\n")

    llm_data = safe_json_loads(raw, fallback={"job_title": "unknown", "seniority": "unknown", "requirements": []})

    if isinstance(llm_data, list):
        llm_data = {"job_title": "unknown", "seniority": "unknown", "requirements": llm_data}
    if not isinstance(llm_data, dict):
        llm_data = {"job_title": "unknown", "seniority": "unknown", "requirements": []}

    llm_data["requirements"] = clean_requirements(
                                                    llm_data.get("requirements") or [],
                                                    max_requirements,
                                                    model_id=model_id,
                                                    load_mode=load_mode,
                                                    fallback_model_id=fallback_model_id,
                                                    aux_model_id=aux_model_id,
                                                    allow_llm=True,
                                                )
    llm_data["_source"] = "llm_few_shot"

    weak_llm = is_weak_requirement_output(
                                            llm_data,
                                            max_requirements=max_requirements,
                                            model_id=model_id,
                                            load_mode=load_mode,
                                            fallback_model_id=fallback_model_id,
                                            aux_model_id=aux_model_id,
                                            allow_llm=True,
                                            )
    fallback_data = fallback_extract_requirements_from_text(job_text, max_requirements)

    hybrid_data = build_hybrid_requirement_result(
                                                    llm_data=llm_data,
                                                    fallback_data=fallback_data,
                                                    max_requirements=max_requirements,
                                                )

    if not isinstance(hybrid_data, dict):
        hybrid_data = llm_data

    hybrid_data.setdefault("_meta", {})
    if isinstance(hybrid_data["_meta"], dict):
        hybrid_data["_meta"]["weak_llm"] = weak_llm
        hybrid_data["_meta"]["prompt_mode"] = "few_shot_langchain"

    hybrid_data["requirements"] = clean_requirements(
                                                        hybrid_data.get("requirements") or [],
                                                        max_requirements,
                                                        model_id=model_id,
                                                        load_mode=load_mode,
                                                        fallback_model_id=fallback_model_id,
                                                        aux_model_id=aux_model_id,
                                                        allow_llm=True,
                                                    )

    print("\nREQUIREMENT EXTRACTION SUMMARY")
    print(json.dumps(hybrid_data.get("_meta", {}), ensure_ascii=False, indent=2))
    print(f"Source: {hybrid_data.get('_source', 'unknown')}")
    print("END REQUIREMENT EXTRACTION SUMMARY\n")

    return hybrid_data

def extract_candidate_summary(
                            cv_text: str,
                            model_id: str,
                            load_mode: str,
                            fallback_model_id: str,
) -> Dict[str, Any]:
    final_user = (
                    "Zo zivotopisu extrahuj anonymizovany profil kandidata.\n\n"
                    "Vrat presne tento JSON tvar:\n"
                    + CANDIDATE_SUMMARY_SCHEMA
                    + "\n\nPravidla:\n"
                    "- nepouzivaj meno, adresu, vek, pohlavie, rodinny stav ani fotku\n"
                    "- uvadzaj len veci, ktore su v CV\n"
                    "- vrat iba JSON, bez komentara\n\n"
                    "CV TEXT:\n"
                    + trim_text(cv_text, 22000)
                    )

    messages = build_lc_messages_from_fewshot(
                                                system_text=SYSTEM_JSON,
                                                examples=CANDIDATE_FEWSHOT,
                                                example_human_template="Zo zivotopisu extrahuj anonymizovany profil kandidata.\n\nCV TEXT:\n{input}",
                                                final_human_input=final_user,
                                                )

    raw = chat_generate_messages(
                                messages,
                                model_id,
                                load_mode,
                                fallback_model_id,
                                max_new_tokens=900,
                                do_sample=False,
                                )

    data = safe_json_loads(raw, fallback={"summary": "Nepodarilo sa spolahlivo extrahovat profil.", "skills": []})
    if not isinstance(data, dict):
        data = {"summary": "Nepodarilo sa spolahlivo extrahovat profil.", "skills": []}

    for key in ["skills", "experience", "education", "languages", "certifications", "risks_or_missing_info"]:
        vals = data.get(key)
        data[key] = dedupe_text_list(vals if isinstance(vals, list) else [], limit=12)

    return data

def evaluate_one_requirement(
                            requirement: Dict[str, Any],
                            evidence: List[str],
                            model_id: str,
                            load_mode: str,
                            fallback_model_id: str,
) -> Dict[str, Any]:
    final_user = (
        "Vyhodnot, ci kandidat splna jednu poziadavku z pracovneho inzeratu.\n"
        "Pouzi iba dodane odkazy z CV. Ak dokaz chyba, nehadaj.\n\n"
        "Vrat presne tento JSON tvar:\n"
        + json.dumps(REQUIREMENT_EVAL_SCHEMA, ensure_ascii=False, indent=2)
        + "\n\nSkorovanie:\n"
        "- 90-100 = jasne splnene\n"
        "- 60-89 = skor splnene alebo ciastocne\n"
        "- 30-59 = slabe/nepriame odkazy\n"
        "- 0-29 = nesplnene alebo chyba dokaz\n\n"
        "POZIADAVKA:\n"
        + json.dumps(requirement, ensure_ascii=False, indent=2)
        + "\n\nODKAZY Z CV:\n"
        + json.dumps(evidence, ensure_ascii=False, indent=2)
    )

    eval_examples = []
    for ex in REQUIREMENT_EVAL_FEWSHOT:
        eval_examples.append(
            {
                "input": (
                    "POZIADAVKA:\n"
                    + json.dumps(ex["input"]["requirement"], ensure_ascii=False, indent=2)
                    + "\n\nODKAZY Z CV:\n"
                    + json.dumps(ex["input"]["evidence"], ensure_ascii=False, indent=2)
                ),
                "output": ex["output"],
            }
        )

    messages = build_lc_messages_from_fewshot(
                                                system_text=SYSTEM_JSON,
                                                examples=eval_examples,
                                                example_human_template="Vyhodnot, ci kandidat splna jednu poziadavku z pracovneho inzeratu.\n\n{input}",
                                                final_human_input=final_user,
                                                )

    raw = chat_generate_messages(
                                messages,
                                model_id,
                                load_mode,
                                fallback_model_id,
                                max_new_tokens=700,
                                do_sample=False,
                                )

    data = safe_json_loads(raw, fallback={})
    if not isinstance(data, dict):
        data = {}

    status = str(data.get("status") or "nejasne")
    if status not in {"splnene", "ciastocne_splnene", "nesplnene", "nejasne"}:
        status = "nejasne"

    try:
        score = float(data.get("score", 0))
    except Exception:
        score = 0.0

    try:
        confidence = float(data.get("confidence", 0))
    except Exception:
        confidence = 0.0

    evidence_used = data.get("evidence_used") if isinstance(data.get("evidence_used"), list) else evidence[:2]
    evidence_used = dedupe_text_list([str(x) for x in evidence_used], limit=3)

    return {
            "requirement_id": str(data.get("requirement_id") or requirement.get("id", "")),
            "requirement": str(data.get("requirement") or requirement.get("text", "")),
            "category": requirement.get("category", "other"),
            "priority": requirement.get("priority", "unknown"),
            "weight": requirement.get("weight", 1.0),
            "status": status,
            "score": max(0.0, min(100.0, score)),
            "confidence": max(0.0, min(1.0, confidence)),
            "evidence_used": evidence_used,
            "explanation": str(data.get("explanation") or "Bez vysvetlenia."),
            "risk_note": str(data.get("risk_note") or ""),
            }

# -----------------------------------------------------------------------------
# 10. REPORTING
# -----------------------------------------------------------------------------

def status_icon(status: str) -> str:
    return {"splnene": "✅","ciastocne_splnene": "🟡","nesplnene": "❌","nejasne": "⚪"}.get(status, "⚪")

def verdict(score: float) -> str:
    if score >= 80:
        return "Vhodny kandidat"
    if score >= 60:
        return "Skor vhodny kandidat"
    if score >= 40:
        return "Ciastocne vhodny / vyzaduje manualne posudenie"
    return "Slaba zhoda s poziciou"

def render_markdown_report(job_data: Dict[str, Any], candidate: Dict[str, Any], evals: List[Dict[str, Any]]) -> str:
    score = weighted_average(evals)
    lines = []
    lines.append("# AI CV validator")
    lines.append("")
    lines.append(f"**Pozicia:** {job_data.get('job_title', 'unknown')}")
    lines.append(f"**Seniorita:** {job_data.get('seniority', 'unknown')}")
    lines.append(f"**Zdroj poziadaviek:** {job_data.get('_source', 'unknown')}")
    lines.append(f"**Celkove skore:** {score:.2f} / 100")
    lines.append(f"**Odporucanie:** {verdict(score)}")
    lines.append("")
    lines.append("> Vystup je odporucanie pre cloveka, nie automaticke rozhodnutie o kandidatovi.")
    lines.append("")

    if candidate:
        lines.append("## Anonymizovany profil kandidata")
        lines.append(str(candidate.get("summary", "")))
        for key, title in [
            ("skills", "Zrucnosti"),
            ("experience", "Skusenosti"),
            ("education", "Vzdelanie"),
            ("languages", "Jazyky"),
            ("certifications", "Certifikaty"),
            ("risks_or_missing_info", "Rizika / chybajuce info"),
        ]:
            vals = candidate.get(key)
            if isinstance(vals, list) and vals:
                lines.append(f"\n**{title}:**")
                for v in vals[:12]:
                    lines.append(f"- {v}")
        lines.append("")

    lines.append("## Vyhodnotenie poziadaviek")
    lines.append("")
    lines.append("| Stav | Poziadavka | Skore | Priorita | Vysvetlenie |")
    lines.append("|---|---|---:|---|---|")
    for r in evals:
        exp = str(r.get("explanation", "")).replace("|", "/")
        req = str(r.get("requirement", "")).replace("|", "/")
        lines.append(
            f"| {status_icon(r.get('status', ''))} {r.get('status', '')} "
            f"| {req} | {float(r.get('score', 0)):.0f} "
            f"| {r.get('priority', '')} / w={r.get('weight', 1)} | {exp} |"
        )

    lines.append("")
    lines.append("## Odkazy a poznamky")
    for r in evals:
        lines.append(f"\n### {status_icon(r.get('status', ''))} {r.get('requirement_id', '')}: {r.get('requirement', '')}")
        if r.get("risk_note"):
            lines.append(f"**Riziko/neistota:** {r.get('risk_note')}")
        ev = r.get("evidence_used") or []
        if isinstance(ev, list) and ev:
            lines.append("**Pouzite odkazy:**")
            for e in ev[:3]:
                short = normalize_space(str(e))[:700]
                lines.append(f"- {short}")

    return "\n".join(lines)

# -----------------------------------------------------------------------------
# 11. MAIN PIPELINE
# -----------------------------------------------------------------------------

def run_validation(
                    cv_file: str,
                    job_url: str,
                    job_text_manual: str,
                    job_schema_xlsx_path: str,
                    manual_position_name: str,
                    model_id: str,
                    fallback_model_id: str,
                    aux_model_id: str,
                    load_mode: str,
                    embed_model_id: str,
                    top_k: int,
                    max_requirements: int,
                    include_candidate_summary: bool,
) -> Tuple[str, str, str]:
    if not cv_file:
        raise gr.Error("Nahraj CV subor.")

    runtime = []
    runtime.append(cuda_summary())
    runtime.append(f"LLM model: {model_id}")
    runtime.append(f"Fallback model: {fallback_model_id}")
    runtime.append(f"Aux model: {aux_model_id or model_id}")
    runtime.append(f"Load mode: {load_mode}")
    runtime.append(f"Embedding model: {embed_model_id}")
    runtime.append(f"Top-K: {top_k}")
    runtime.append(f"Max requirements: {max_requirements}")
    runtime.append(f"Schema XLSX path: {job_schema_xlsx_path or DEFAULT_JOB_SCHEMA_XLSX_PATH}")
    runtime.append(f"Manual position: {manual_position_name or '-'}")
    runtime.append("Prompt mode: LangChain few-shot")

    cv_text = load_document(cv_file)
    if len(cv_text) < 100:
        raise gr.Error("Z CV sa podarilo vytiahnut velmi malo textu. Skus iny format, idealne PDF/DOCX.")
    runtime.append(f"CV text: {len(cv_text):,} znakov")

    job_requirement_schema_text, prompt_schema_source = load_job_requirement_schema_text(
        job_schema_xlsx_path or DEFAULT_JOB_SCHEMA_XLSX_PATH
    )
    runtime.append(f"Prompt schema source: {prompt_schema_source}")

    _, _, model_info = load_llm(model_id, load_mode, fallback_model_id)
    runtime.append(model_info)

    if manual_position_name and manual_position_name.strip():
        job_data = load_manual_job_requirements_from_excel(
                                                            position_query=manual_position_name,
                                                            schema_xlsx_path=job_schema_xlsx_path or DEFAULT_JOB_SCHEMA_XLSX_PATH,
                                                            max_requirements=max_requirements,
                                                            model_id=model_id,
                                                            load_mode=load_mode,
                                                            fallback_model_id=fallback_model_id,
                                                            aux_model_id=aux_model_id,
                                                            )
        runtime.append("Zdroj pozicie: manualny schema XLSX katalog")
    else:
        job_text = ""
        if job_text_manual and job_text_manual.strip():
            job_text = normalize_space(job_text_manual)
            runtime.append("Inzerat: pouzity manualne vlozeny text")
        elif job_url and job_url.strip():
            job_text = scrape_url(job_url)
            runtime.append(f"Inzerat: nacitany z URL, {len(job_text):,} znakov")
        else:
            raise gr.Error("Zadaj URL inzeratu, vloz text inzeratu manualne, alebo vypln manualnu poziciu zo schema XLSX.")

        if len(job_text) < 100:
            raise gr.Error("Z inzeratu sa podarilo ziskat velmi malo textu. Vloz text inzeratu manualne.")

        job_data = extract_job_requirements(
                                            job_text,
                                            job_requirement_schema_text,
                                            model_id,
                                            load_mode,
                                            fallback_model_id,
                                            aux_model_id,
                                            max_requirements,
                                            )
    requirements = job_data.get("requirements", [])
    if not requirements:
        raise gr.Error("LLM ani fallback neextrahovali ziadne poziadavky z inzeratu. Skus vlozit cistejsi text inzeratu.")
    runtime.append(f"Extrahovane poziadavky: {len(requirements)}")
    runtime.append(f"Zdroj poziadaviek: {job_data.get('_source', 'unknown')}")

    meta = job_data.get("_meta", {})
    if isinstance(meta, dict):
        runtime.append(f"LLM count: {meta.get('llm_count', 0)}")
        runtime.append(f"Fallback count: {meta.get('fallback_count', 0)}")
        runtime.append(f"Weak LLM: {meta.get('weak_llm', False)}")
        runtime.append(f"Merged count: {meta.get('merged_count', 0)}")
        runtime.append(f"Prompt mode meta: {meta.get('prompt_mode', 'few_shot_langchain')}")

    candidate = {}
    if include_candidate_summary:
        candidate = extract_candidate_summary(cv_text, model_id, load_mode, fallback_model_id)

    chunks = chunk_text(cv_text, CHUNK_WORDS, CHUNK_OVERLAP)
    index, _ = build_faiss_index(chunks, embed_model_id)
    runtime.append(f"CV chunks: {len(chunks)}")

    evals = []
    for req in requirements:
        evidence = rag_search(req.get("text", ""), chunks, index, embed_model_id, int(top_k))
        ev = evaluate_one_requirement(req, evidence, model_id, load_mode, fallback_model_id)
        evals.append(ev)

    final = {
            "job": job_data,
            "candidate_profile": candidate,
            "overall_score": weighted_average(evals),
            "verdict": verdict(weighted_average(evals)),
            "evaluations": evals,
            "runtime": runtime,
            }

    md = render_markdown_report(job_data, candidate, evals)
    js = json.dumps(final, ensure_ascii=False, indent=2)
    return md, js, "\n".join(runtime)

def gradio_run_wrapper(*args):
    try:
        return run_validation(*args)
    except gr.Error:
        raise
    except Exception as exc:
        err = f"Chyba: {type(exc).__name__}: {exc}\n\n{traceback.format_exc()}"
        return "# Chyba pri spracovani\n\n```text\n" + err + "\n```", "{}", err

# -----------------------------------------------------------------------------
# 12. GRADIO UI
# -----------------------------------------------------------------------------

def build_ui():
    with gr.Blocks(title="Lokalny AI CV Validator") as demo:
        gr.Markdown("")

        with gr.Row():
            with gr.Column(scale=1):
                cv_file = gr.File(
                    label="CV subor",
                    file_types=[".pdf", ".docx", ".doc", ".rtf", ".txt", ".md"],
                    type="filepath",
                )
                job_url = gr.Textbox(label="URL inzeratu", placeholder="https://")
                job_text_manual = gr.Textbox(
                    label="Alebo vloz text inzeratu manualne",
                    lines=8,
                    placeholder="Text pracovnej ponuky",
                )

                with gr.Accordion("Externy schema XLSX", open=False):
                    job_schema_xlsx_path = gr.Textbox(
                        label="Schema XLSX path",
                        value=DEFAULT_JOB_SCHEMA_XLSX_PATH,
                    )
                    manual_position_name = gr.Textbox(
                        label="Manualna pozicia zo schema XLSX",
                        placeholder="napr. python_backend_medior alebo Python developer",
                    )

                with gr.Accordion("Model nastavenia", open=False):
                    model_id = gr.Textbox(label="LLM model z Hugging Face", value=DEFAULT_LLM_MODEL_ID)
                    fallback_model_id = gr.Textbox(label="Fallback LLM model", value=DEFAULT_FALLBACK_LLM_MODEL_ID)
                    aux_model_id = gr.Textbox(label="Aux LLM model pre canonicalizaciu/genericnost", value=DEFAULT_AUX_LLM_MODEL_ID or DEFAULT_LLM_MODEL_ID)
                    load_mode = gr.Dropdown(
                        label="Load mode",
                        choices=["auto", "bnb_4bit", "fp16_gpu", "cpu"],
                        value=LLM_LOAD_MODE,
                    )
                    embed_model_id = gr.Textbox(label="Embedding model z Hugging Face", value=DEFAULT_EMBED_MODEL_ID)
                    top_k = gr.Slider(label="Top-K dokazov z CV", minimum=2, maximum=20, step=1, value=DEFAULT_TOP_K)
                    max_requirements = gr.Slider(
                        label="Max pocet poziadaviek z inzeratu",
                        minimum=5,
                        maximum=25,
                        step=1,
                        value=DEFAULT_MAX_REQUIREMENTS,
                    )
                    include_candidate_summary = gr.Checkbox(label="Extrahovat anonymizovany profil kandidata", value=True)

                with gr.Row():
                    run_btn = gr.Button("Spustit validaciu", variant="primary")
                    unload_btn = gr.Button("Uvolnit model z VRAM")

            with gr.Column(scale=2):
                report_md = gr.Markdown(label="Report")
                runtime_info = gr.Textbox(label="Runtime info", lines=8)
                json_report = gr.Code(label="JSON report", language="json", lines=20)

        run_btn.click(
                        fn=gradio_run_wrapper,
                        inputs=[
                                cv_file,
                                job_url,
                                job_text_manual,
                                job_schema_xlsx_path,
                                manual_position_name,
                                model_id,
                                fallback_model_id,
                                aux_model_id,
                                load_mode,
                                embed_model_id,
                                top_k,
                                max_requirements,
                                include_candidate_summary,
                                ],
                        outputs=[report_md, json_report, runtime_info],
                        )
        unload_btn.click(fn=unload_llm, inputs=[], outputs=[runtime_info])

    return demo

if __name__ == "__main__":
    demo = build_ui()
    demo.launch(server_name="127.0.0.1", server_port=7860, inbrowser=True)
